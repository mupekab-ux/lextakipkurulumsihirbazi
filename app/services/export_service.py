# -*- coding: utf-8 -*-
"""
Dışa aktarma ve yedekleme işlemleri.
"""

import os
import shutil
from services.base import *

try:
    import pandas as pd
    from openpyxl import Workbook
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor, Pt, Mm
    from docx.enum.section import WD_ORIENT, WD_SECTION
except ImportError:
    pd = None
    Workbook = None
    Document = None

__all__ = [
    "EXPORT_FIELDS",
    "EXPORT_HEADERS",
    "DATE_FIELDS",
    "HEADER_LABELS",
    "STATUS_COLOR_FIELDS",
    "REQUIRED_DB_TABLES",
    "export_dosyalar_to_csv",
    "export_dosyalar_to_xlsx",
    "export_dosyalar_to_docx",
    "backup_database",
    "validate_database_file",
]

# Dışa aktarma alan tanımları
EXPORT_FIELDS = [
    "buro_takip_no",
    "dosya_esas_no",
    "muvekkil_adi",
    "muvekkil_rolu",
    "karsi_taraf",
    "dosya_konusu",
    "mahkeme_adi",
    "dava_acilis_tarihi",
    "durusma_tarihi",
    "dava_durumu",
    "is_tarihi",
    "aciklama",
    "tekrar_dava_durumu_2",
    "is_tarihi_2",
    "aciklama_2",
]

EXPORT_HEADERS: dict[str, str] = {
    "buro_takip_no": "Büro Takip No",
    "dosya_esas_no": "Dosya Esas No",
    "muvekkil_adi": "Müvekkil Adı",
    "muvekkil_rolu": "Müvekkil Rolü",
    "karsi_taraf": "Karşı Taraf",
    "dosya_konusu": "Dosya Konusu",
    "mahkeme_adi": "Mahkeme Adı",
    "dava_acilis_tarihi": "Dava Açılış Tarihi",
    "durusma_tarihi": "Duruşma Tarihi",
    "dava_durumu": "Dava Durumu",
    "is_tarihi": "İş Tarihi",
    "aciklama": "Açıklama",
    "tekrar_dava_durumu_2": "Dava Durumu 2",
    "is_tarihi_2": "İş Tarihi 2",
    "aciklama_2": "Açıklama 2",
}

DATE_FIELDS = {"dava_acilis_tarihi", "durusma_tarihi", "is_tarihi", "is_tarihi_2"}

HEADER_LABELS = [EXPORT_HEADERS.get(field, field) for field in EXPORT_FIELDS]

STATUS_COLOR_FIELDS = [
    ("dava_durumu_color", EXPORT_FIELDS.index("dava_durumu")),
    ("tekrar_dava_durumu_2_color", EXPORT_FIELDS.index("tekrar_dava_durumu_2")),
]

# Gerekli veritabanı tabloları
REQUIRED_DB_TABLES = {"dosyalar", "users", "statuses"}


def _format_export_value(field: str, value: Any) -> Any:
    """Dışa aktarma için değeri formatlar."""
    if value is None:
        return ""
    if field in DATE_FIELDS and value:
        return iso_to_tr(value)
    return value


def _prepare_export_row(row: Dict[str, Any]) -> List[Any]:
    """Dışa aktarma için satırı hazırlar."""
    return [_format_export_value(key, row.get(key)) for key in EXPORT_FIELDS]


def _prepare_export_dict(row: Dict[str, Any]) -> Dict[str, Any]:
    """Dışa aktarma için sözlük formatında satır hazırlar."""
    values = _prepare_export_row(row)
    return dict(zip(HEADER_LABELS, values))


def _calculate_column_weights(prepared_rows: List[List[Any]]) -> List[float]:
    """Sütun genişliklerini hesaplar."""
    if not prepared_rows:
        return [1.0] * len(EXPORT_FIELDS)

    col_count = len(EXPORT_FIELDS)
    max_lengths = [len(str(HEADER_LABELS[i])) for i in range(col_count)]

    for row in prepared_rows:
        for i, value in enumerate(row):
            text_len = len(str(value) if value else "")
            if text_len > max_lengths[i]:
                max_lengths[i] = text_len

    total = sum(max_lengths) or 1
    return [length / total for length in max_lengths]


def _scale_widths(weights: List[float], available_width: float) -> List[float]:
    """Ağırlıklara göre genişlikleri ölçekler."""
    total_weight = sum(weights) or 1.0
    return [(w / total_weight) * available_width for w in weights]


def _apply_docx_cell_colors(cell, bg_hex: str | None) -> None:
    """Word belgesi hücresine arka plan rengi uygular."""
    normalized_bg = normalize_hex(bg_hex)
    if not normalized_bg:
        return
    text_color = get_status_text_color(normalized_bg)
    normalized_text = normalize_hex(text_color) or "000000"

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), normalized_bg)
    shading.set(qn("w:color"), normalized_text)

    rgb = RGBColor.from_string(normalized_text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = rgb


def export_dosyalar_to_csv(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini CSV olarak dışa aktarır."""
    if pd is None:
        raise ImportError("pandas modülü yüklü değil")
    data = [_prepare_export_dict(row) for row in rows]
    df = pd.DataFrame(data, columns=HEADER_LABELS)
    df.to_csv(path, index=False, encoding="utf-8")


def export_dosyalar_to_xlsx(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini XLSX olarak dışa aktarır."""
    if Workbook is None:
        raise ImportError("openpyxl modülü yüklü değil")
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Dosyalar"

    # Başlık satırı stilleri
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )

    # Başlık satırını ekle
    ws.append(HEADER_LABELS)
    for col_idx, _ in enumerate(HEADER_LABELS, 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Veri satırları stilleri
    data_font = Font(name="Calibri", size=10)
    data_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Veri satırlarını ekle
    for row_idx, row in enumerate(rows, 2):
        prepared = _prepare_export_row(row)
        ws.append(prepared)

        # Satır stillerini uygula
        for col_idx, value in enumerate(prepared, 1):
            cell = ws.cell(row=row_idx, column=col_idx)
            cell.font = data_font
            cell.alignment = data_alignment
            cell.border = thin_border

        # Dava durumu hücrelerine renk uygula
        for color_field, column_index in STATUS_COLOR_FIELDS:
            color_hex = normalize_hex(row.get(color_field))
            if color_hex:
                cell = ws.cell(row=row_idx, column=column_index + 1)
                cell.fill = PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")
                text_color = get_status_text_color(color_hex).replace("#", "")
                cell.font = Font(name="Calibri", size=10, color=text_color)

    # Sütun genişliklerini ayarla
    column_widths = {
        1: 12,   # Büro Takip No
        2: 18,   # Dosya Esas No
        3: 20,   # Müvekkil Adı
        4: 12,   # Müvekkil Rolü
        5: 20,   # Karşı Taraf
        6: 25,   # Dosya Konusu
        7: 20,   # Mahkeme Adı
        8: 14,   # Dava Açılış Tarihi
        9: 14,   # Duruşma Tarihi
        10: 25,  # Dava Durumu
        11: 14,  # İş Tarihi
        12: 30,  # Açıklama
        13: 25,  # Dava Durumu 2
        14: 14,  # İş Tarihi 2
        15: 30,  # Açıklama 2
    }
    for col_idx, width in column_widths.items():
        if col_idx <= len(HEADER_LABELS):
            ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Satır yüksekliğini ayarla
    ws.row_dimensions[1].height = 30  # Başlık satırı
    for row_idx in range(2, len(rows) + 2):
        ws.row_dimensions[row_idx].height = 20

    # Filtre ekle
    ws.auto_filter.ref = ws.dimensions

    # Sayfa yazdırma ayarları
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0

    wb.save(path)


def export_dosyalar_to_docx(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini Word belgesi olarak dışa aktarır."""
    if Document is None:
        raise ImportError("python-docx modülü yüklü değil")
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Normal stil ayarları
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(8)

    # Sayfa ayarları - Yatay (Landscape) ve geniş marjlar
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width, section.page_height = new_width, new_height
        # Daha dar marjlar
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        section.top_margin = Mm(15)
        section.bottom_margin = Mm(15)

    # Başlık ekle
    title = doc.add_paragraph()
    title_run = title.add_run("Dosya Listesi")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Calibri"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tarih bilgisi
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_run.font.size = Pt(9)
    date_run.font.name = "Calibri"
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    prepared_rows = [_prepare_export_row(row) for row in rows]

    # Tablo oluştur
    table = doc.add_table(rows=1, cols=len(HEADER_LABELS))
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Başlık hücreleri
    header_cells = table.rows[0].cells
    for index, header in enumerate(HEADER_LABELS):
        cell = header_cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        run = paragraph.add_run(str(header))
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        # Başlık arka plan rengi
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:fill"), "2c3e50")
        tcPr.append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Veri satırları
    for row, prepared_row in zip(rows, prepared_rows):
        row_cells = table.add_row().cells
        for index, value in enumerate(prepared_row):
            text = "" if value is None else str(value)
            # Çok uzun metinleri kısalt
            if len(text) > 100:
                text = text[:97] + "..."
            cell = row_cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            lines = text.split("\n")
            for idx, part in enumerate(lines):
                run = paragraph.add_run(part)
                run.font.size = Pt(8)
                run.font.name = "Calibri"
                if idx < len(lines) - 1:
                    run.add_break()

        # Dava durumu renklerini uygula
        for field, column_index in STATUS_COLOR_FIELDS:
            color_hex = normalize_hex(row.get(field))
            if not color_hex:
                continue
            _apply_docx_cell_colors(row_cells[column_index], color_hex)

    # Sabit sütun genişlikleri (mm cinsinden)
    column_widths_mm = [
        12,   # Büro Takip No
        18,   # Dosya Esas No
        18,   # Müvekkil Adı
        10,   # Müvekkil Rolü
        18,   # Karşı Taraf
        22,   # Dosya Konusu
        18,   # Mahkeme Adı
        12,   # Dava Açılış Tarihi
        12,   # Duruşma Tarihi
        22,   # Dava Durumu
        12,   # İş Tarihi
        25,   # Açıklama
        22,   # Dava Durumu 2
        12,   # İş Tarihi 2
        25,   # Açıklama 2
    ]

    for col_index, width_mm in enumerate(column_widths_mm):
        if col_index < len(HEADER_LABELS):
            table.columns[col_index].width = Mm(width_mm)
            for row in table.rows:
                row.cells[col_index].width = Mm(width_mm)

    # Toplam kayıt sayısı
    footer = doc.add_paragraph()
    footer_run = footer.add_run(f"\nToplam: {len(rows)} kayıt")
    footer_run.font.size = Pt(9)
    footer_run.font.name = "Calibri"
    footer_run.italic = True

    doc.save(path)


def backup_database(dest_path: str) -> None:
    """Veritabanı dosyasını belirtilen yola kopyalar."""
    shutil.copy2(DB_PATH, dest_path)


def validate_database_file(
    path: str, required_tables: Iterable[str] | None = None
) -> tuple[bool, Set[str]]:
    """Seçilen veritabanı dosyasının gerekli tabloları içerip içermediğini kontrol eder."""
    required = {table.lower() for table in (required_tables or REQUIRED_DB_TABLES)}
    if not path or not os.path.exists(path):
        return False, required

    try:
        conn = sqlite3.connect(path)
        cur = conn.cursor()
        cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = {row[0].lower() for row in cur.fetchall()}
    except sqlite3.Error:
        return False, required
    finally:
        try:
            conn.close()
        except Exception:
            pass

    missing = required - tables
    return len(missing) == 0, missing
