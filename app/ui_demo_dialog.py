# -*- coding: utf-8 -*-
"""
TakibiEsasi Demo Modu UI Bileşenleri

- DemoRegistrationDialog: İlk açılışta e-posta kayıt dialogu
- DemoStatusWidget: Ana pencerede demo durumu banner'ı
- DemoExpiredDialog: Demo süresi dolduğunda gösterilen dialog
"""

from PyQt6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel,
    QLineEdit, QPushButton, QMessageBox, QFrame,
    QSpacerItem, QSizePolicy, QWidget
)
from PyQt6.QtCore import Qt, pyqtSignal
from PyQt6.QtGui import QFont, QPixmap, QPainter, QColor, QIcon
import re
import webbrowser


class DemoRegistrationDialog(QDialog):
    """
    Demo başlatmak için e-posta kayıt dialogu.
    İlk açılışta gösterilir.
    """

    demo_started = pyqtSignal(str)  # e-posta ile sinyal

    def __init__(self, parent=None):
        super().__init__(parent)
        self.email = None
        self.result_action = None  # "demo", "license", "cancel"
        self._setup_ui()

    def _setup_ui(self):
        self.setWindowTitle("TakibiEsasi - Hoş Geldiniz")
        self.setFixedSize(480, 380)
        self.setModal(True)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)

        # Ana layout
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(50, 40, 50, 40)

        # Logo/İkon alanı
        icon_label = QLabel()
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        icon_label.setText("⚖️")
        icon_label.setStyleSheet("font-size: 48px;")
        layout.addWidget(icon_label)

        # Başlık
        title = QLabel("14 Gün Ücretsiz Deneyin")
        title.setFont(QFont("Segoe UI", 20, QFont.Weight.Bold))
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("color: #FBBF24; margin-bottom: 5px;")
        layout.addWidget(title)

        # Alt başlık
        subtitle = QLabel("Tüm özelliklere tam erişim\nKredi kartı gerekmez")
        subtitle.setAlignment(Qt.AlignmentFlag.AlignCenter)
        subtitle.setStyleSheet("color: #888; font-size: 13px; line-height: 1.4;")
        layout.addWidget(subtitle)

        layout.addSpacing(10)

        # E-posta input
        self.email_input = QLineEdit()
        self.email_input.setPlaceholderText("E-posta adresiniz")
        self.email_input.setMinimumHeight(50)
        self.email_input.setStyleSheet("""
            QLineEdit {
                border: 2px solid #333;
                border-radius: 10px;
                padding: 12px 18px;
                font-size: 15px;
                background: #1a1a1f;
                color: #fff;
            }
            QLineEdit:focus {
                border-color: #FBBF24;
            }
            QLineEdit::placeholder {
                color: #666;
            }
        """)
        self.email_input.returnPressed.connect(self._on_start_demo)
        layout.addWidget(self.email_input)

        # Bilgi notu
        info_label = QLabel("✓ Verileriniz satın alınca korunur")
        info_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        info_label.setStyleSheet("color: #10B981; font-size: 12px;")
        layout.addWidget(info_label)

        layout.addSpacing(5)

        # Demo başlat butonu
        self.start_btn = QPushButton("Demo'yu Başlat")
        self.start_btn.setMinimumHeight(55)
        self.start_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.start_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #FBBF24, stop:1 #D97706);
                color: #000;
                border: none;
                border-radius: 10px;
                font-size: 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #FCD34D, stop:1 #F59E0B);
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #D97706, stop:1 #B45309);
            }
        """)
        self.start_btn.clicked.connect(self._on_start_demo)
        layout.addWidget(self.start_btn)

        layout.addSpacing(10)

        # Lisans linki
        license_link = QLabel('<a href="#" style="color: #FBBF24; text-decoration: none;">Zaten lisans anahtarınız var mı?</a>')
        license_link.setAlignment(Qt.AlignmentFlag.AlignCenter)
        license_link.setOpenExternalLinks(False)
        license_link.linkActivated.connect(self._on_have_license)
        layout.addWidget(license_link)

        # Arka plan stili
        self.setStyleSheet("""
            QDialog {
                background-color: #0f0f14;
            }
        """)

    def _on_start_demo(self):
        """Demo başlat butonuna tıklandı"""
        email = self.email_input.text().strip()

        # E-posta validasyonu
        if not email:
            QMessageBox.warning(
                self,
                "E-posta Gerekli",
                "Lütfen e-posta adresinizi girin."
            )
            self.email_input.setFocus()
            return

        # Basit e-posta format kontrolü
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_pattern, email):
            QMessageBox.warning(
                self,
                "Geçersiz E-posta",
                "Lütfen geçerli bir e-posta adresi girin."
            )
            self.email_input.setFocus()
            self.email_input.selectAll()
            return

        self.email = email
        self.result_action = "demo"
        self.accept()

    def _on_have_license(self):
        """Lisans var linkine tıklandı"""
        self.result_action = "license"
        self.email = None
        self.accept()

    def get_result(self):
        """Dialog sonucunu al"""
        return {
            "action": self.result_action,
            "email": self.email
        }


class DemoStatusWidget(QFrame):
    """
    Ana pencerede gösterilecek demo durum banner'ı.
    Kalan süreyi ve satın alma butonunu gösterir.
    """

    buy_clicked = pyqtSignal()
    license_clicked = pyqtSignal()

    def __init__(self, days_remaining: int, parent=None):
        super().__init__(parent)
        self.days_remaining = days_remaining
        self._setup_ui()

    def _setup_ui(self):
        # Kalan süreye göre stil
        if self.days_remaining > 7:
            bg_color = "rgba(251, 191, 36, 0.1)"
            border_color = "rgba(251, 191, 36, 0.3)"
            text_color = "#FBBF24"
            icon = "🎯"
            text = f"Demo Sürümü - {self.days_remaining} gün kaldı"
        elif self.days_remaining > 3:
            bg_color = "rgba(245, 158, 11, 0.15)"
            border_color = "rgba(245, 158, 11, 0.4)"
            text_color = "#F59E0B"
            icon = "⏳"
            text = f"Demo Sürümü - {self.days_remaining} gün kaldı"
        elif self.days_remaining > 0:
            bg_color = "rgba(239, 68, 68, 0.15)"
            border_color = "rgba(239, 68, 68, 0.4)"
            text_color = "#EF4444"
            icon = "⚠️"
            text = f"Demo Sürümü - Son {self.days_remaining} gün!"
        else:
            bg_color = "rgba(239, 68, 68, 0.2)"
            border_color = "rgba(239, 68, 68, 0.5)"
            text_color = "#EF4444"
            icon = "❌"
            text = "Demo Süresi Doldu"

        self.setStyleSheet(f"""
            QFrame {{
                background: {bg_color};
                border: 1px solid {border_color};
                border-radius: 8px;
            }}
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(15, 10, 15, 10)
        layout.setSpacing(15)

        # İkon ve metin
        label = QLabel(f"{icon} {text}")
        label.setStyleSheet(f"color: {text_color}; font-weight: bold; font-size: 13px; border: none; background: transparent;")
        layout.addWidget(label)

        layout.addStretch()

        # Lisans gir butonu
        license_btn = QPushButton("Lisans Gir")
        license_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        license_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #FBBF24;
                border: 1px solid #FBBF24;
                border-radius: 5px;
                padding: 6px 15px;
                font-size: 12px;
            }
            QPushButton:hover {
                background: rgba(251, 191, 36, 0.1);
            }
        """)
        license_btn.clicked.connect(self.license_clicked.emit)
        layout.addWidget(license_btn)

        # Satın al butonu
        buy_btn = QPushButton("Satın Al")
        buy_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        buy_btn.setStyleSheet("""
            QPushButton {
                background: #FBBF24;
                color: #000;
                border: none;
                border-radius: 5px;
                padding: 6px 15px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background: #F59E0B;
            }
        """)
        buy_btn.clicked.connect(self._on_buy_clicked)
        layout.addWidget(buy_btn)

    def _on_buy_clicked(self):
        """Satın al butonuna tıklandı"""
        webbrowser.open("https://takibiesasi.com")
        self.buy_clicked.emit()

    def update_days(self, days_remaining: int):
        """Kalan gün sayısını güncelle"""
        self.days_remaining = days_remaining
        # Widget'ı yeniden oluştur
        for i in reversed(range(self.layout().count())):
            self.layout().itemAt(i).widget().setParent(None)
        self._setup_ui()


class DemoExpiredDialog(QDialog):
    """
    Demo süresi dolduğunda gösterilecek dialog.
    Kullanıcıyı satın almaya veya lisans girmeye yönlendirir.
    """

    # Dialog result codes
    RESULT_BUY = 1
    RESULT_LICENSE = 2
    RESULT_EXIT = 0

    def __init__(self, expired_days_ago: int = 0, parent=None):
        super().__init__(parent)
        self.expired_days_ago = expired_days_ago
        self._setup_ui()

    def _setup_ui(self):
        self.setWindowTitle("Demo Süresi Doldu")
        self.setFixedSize(450, 320)
        self.setModal(True)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)

        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(50, 40, 50, 40)

        # İkon
        icon_label = QLabel("⏰")
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        icon_label.setStyleSheet("font-size: 56px;")
        layout.addWidget(icon_label)

        # Başlık
        title = QLabel("Demo Süreniz Doldu")
        title.setFont(QFont("Segoe UI", 18, QFont.Weight.Bold))
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title.setStyleSheet("color: #EF4444;")
        layout.addWidget(title)

        # Açıklama
        desc_text = "Verileriniz güvende!"
        if self.expired_days_ago > 0:
            desc_text += f"\n({self.expired_days_ago} gün önce doldu)"
        desc_text += "\nSatın alarak kaldığınız yerden devam edin."

        desc = QLabel(desc_text)
        desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
        desc.setStyleSheet("color: #888; font-size: 13px; line-height: 1.5;")
        layout.addWidget(desc)

        layout.addSpacing(10)

        # Butonlar
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(15)

        # Satın al butonu
        buy_btn = QPushButton("Satın Al")
        buy_btn.setMinimumHeight(50)
        buy_btn.setMinimumWidth(150)
        buy_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        buy_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #FBBF24, stop:1 #D97706);
                color: #000;
                border: none;
                border-radius: 8px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:1,
                    stop:0 #FCD34D, stop:1 #F59E0B);
            }
        """)
        buy_btn.clicked.connect(self._on_buy)
        btn_layout.addWidget(buy_btn)

        # Lisans gir butonu
        license_btn = QPushButton("Lisans Gir")
        license_btn.setMinimumHeight(50)
        license_btn.setMinimumWidth(150)
        license_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        license_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #FBBF24;
                border: 2px solid #FBBF24;
                border-radius: 8px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background: rgba(251, 191, 36, 0.1);
            }
        """)
        license_btn.clicked.connect(self._on_license)
        btn_layout.addWidget(license_btn)

        layout.addLayout(btn_layout)

        # Kapat butonu
        close_btn = QPushButton("Çıkış")
        close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        close_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #666;
                border: none;
                font-size: 12px;
            }
            QPushButton:hover {
                color: #888;
            }
        """)
        close_btn.clicked.connect(self._on_exit)
        layout.addWidget(close_btn, alignment=Qt.AlignmentFlag.AlignCenter)

        # Arka plan stili
        self.setStyleSheet("""
            QDialog {
                background-color: #0f0f14;
            }
        """)

    def _on_buy(self):
        """Satın al butonuna tıklandı"""
        webbrowser.open("https://takibiesasi.com")
        self.done(self.RESULT_BUY)

    def _on_license(self):
        """Lisans gir butonuna tıklandı"""
        self.done(self.RESULT_LICENSE)

    def _on_exit(self):
        """Çıkış butonuna tıklandı"""
        self.done(self.RESULT_EXIT)


class DemoWatermark:
    """
    Export edilen belgelere demo watermark eklemek için yardımcı sınıf.
    """

    WATERMARK_TEXT = "DEMO SÜRÜM - takibiesasi.com"

    @staticmethod
    def add_to_pdf(pdf_canvas, page_width: int, page_height: int, opacity: float = 0.3):
        """
        PDF sayfasına watermark ekle.

        Args:
            pdf_canvas: ReportLab canvas objesi
            page_width: Sayfa genişliği
            page_height: Sayfa yüksekliği
            opacity: Opaklık (0-1)
        """
        try:
            from reportlab.lib.colors import Color

            pdf_canvas.saveState()
            pdf_canvas.setFillColor(Color(0.5, 0.5, 0.5, alpha=opacity))
            pdf_canvas.setFont("Helvetica-Bold", 12)

            # Sağ alt köşe
            text_width = pdf_canvas.stringWidth(DemoWatermark.WATERMARK_TEXT, "Helvetica-Bold", 12)
            x = page_width - text_width - 20
            y = 20

            pdf_canvas.drawString(x, y, DemoWatermark.WATERMARK_TEXT)
            pdf_canvas.restoreState()
        except Exception as e:
            print(f"PDF watermark hatası: {e}")

    @staticmethod
    def add_to_excel(worksheet, max_row: int, max_col: int):
        """
        Excel çalışma sayfasına watermark ekle.

        Args:
            worksheet: openpyxl worksheet objesi
            max_row: Maksimum satır
            max_col: Maksimum sütun
        """
        try:
            from openpyxl.styles import Font, Alignment
            from openpyxl.comments import Comment

            # Footer'a ekle
            worksheet.oddFooter.center.text = DemoWatermark.WATERMARK_TEXT
            worksheet.oddFooter.center.font = "Arial,Bold"
            worksheet.oddFooter.center.size = 10

            # İlk hücreye yorum ekle
            comment = Comment(
                f"Bu belge TakibiEsasi Demo sürümü ile oluşturulmuştur.\n{DemoWatermark.WATERMARK_TEXT}",
                "TakibiEsasi"
            )
            worksheet.cell(row=1, column=1).comment = comment

        except Exception as e:
            print(f"Excel watermark hatası: {e}")

    @staticmethod
    def add_to_word(document):
        """
        Word belgesine watermark ekle.

        Args:
            document: python-docx Document objesi
        """
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # Footer'a ekle
            for section in document.sections:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                run = footer_para.add_run(DemoWatermark.WATERMARK_TEXT)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.italic = True

        except Exception as e:
            print(f"Word watermark hatası: {e}")


# Dialog test kodu
if __name__ == "__main__":
    import sys
    from PyQt6.QtWidgets import QApplication

    app = QApplication(sys.argv)

    # Registration dialog test
    # dialog = DemoRegistrationDialog()
    # if dialog.exec() == QDialog.DialogCode.Accepted:
    #     print(f"Result: {dialog.get_result()}")

    # Expired dialog test
    dialog = DemoExpiredDialog(expired_days_ago=3)
    result = dialog.exec()
    print(f"Result code: {result}")

    # Status widget test
    # widget = DemoStatusWidget(days_remaining=5)
    # widget.show()

    sys.exit(app.exec())
