# -*- coding: utf-8 -*-
from typing import Any, Dict, Iterable, List, Optional, Set
import csv
import os
import sqlite3
import shutil
import logging
from pathlib import Path
from datetime import datetime, date, timedelta
from decimal import Decimal, ROUND_HALF_UP, InvalidOperation
import calendar

try:  # pragma: no cover - runtime import guard
    from app.attachments import (
        AttachmentError,
        add_attachments as store_case_attachments,
        delete_attachment as remove_case_attachment,
        export_attachment as export_case_attachment,
        list_attachments as list_case_attachments,
        update_attachment_source as update_case_attachment_source,
    )
except ModuleNotFoundError:  # pragma: no cover
    from attachments import (
        AttachmentError,
        add_attachments as store_case_attachments,
        delete_attachment as remove_case_attachment,
        export_attachment as export_case_attachment,
        list_attachments as list_case_attachments,
        update_attachment_source as update_case_attachment_source,
    )

try:  # pragma: no cover - runtime import guard
    from app.db import (
        get_connection,
        DB_PATH,
        DEFAULT_ROLE_PERMISSIONS,
        PERMISSION_ACTIONS,
        timed_query,
    )
except ModuleNotFoundError:  # pragma: no cover
    from db import (
        get_connection,
        DB_PATH,
        DEFAULT_ROLE_PERMISSIONS,
        PERMISSION_ACTIONS,
        timed_query,
    )
from openpyxl import Workbook
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Mm
from docx.enum.section import WD_ORIENT, WD_SECTION
import pandas as pd
try:  # pragma: no cover - runtime import guard
    from app.utils import (
        hash_password,
        verify_password,
        iso_to_tr,
        tr_to_iso,
        normalize_hex,
        normalize_str,
        resolve_owner_label,
        get_status_text_color,
        USER_ROLE_CHOICES,
        format_tl,
        get_attachments_dir,
        turkish_casefold,
    )
except ModuleNotFoundError:  # pragma: no cover
    from utils import (
        hash_password,
        verify_password,
        iso_to_tr,
        tr_to_iso,
        normalize_hex,
        normalize_str,
        resolve_owner_label,
        get_status_text_color,
        USER_ROLE_CHOICES,
        format_tl,
        get_attachments_dir,
        turkish_casefold,
    )

logger = logging.getLogger(__name__)

from PyQt6.QtCore import QDate


INSTALLMENT_OVERDUE_STATUS = "Gecikmiş"
AUTO_PAYMENT_NOTE = "Taksit ödemesi (otomatik)"
HARICI_AUTO_PAYMENT_NOTE = AUTO_PAYMENT_NOTE


def tl_to_cents(value) -> int:
    """Convert TL value to cents safely."""

    if value in (None, ""):
        return 0
    try:
        if isinstance(value, str):
            cleaned = value.strip()
            if not cleaned:
                return 0
            if "," in cleaned:
                normalized = cleaned.replace(".", "").replace(",", ".")
            else:
                normalized = cleaned.replace(",", "")
            quantized = Decimal(normalized)
        else:
            quantized = Decimal(str(value))
    except (InvalidOperation, ValueError, TypeError):
        return 0
    cents = (quantized * Decimal("100")).quantize(
        Decimal("1"), rounding=ROUND_HALF_UP
    )
    try:
        return int(cents)
    except (ValueError, TypeError):
        return 0


def cents_to_tl(value) -> float:
    """Convert cents value to TL safely."""

    if value in (None, ""):
        return 0.0
    try:
        cents = Decimal(int(value))
    except (ValueError, TypeError, InvalidOperation):
        return 0.0
    tl_amount = cents / Decimal("100")
    return float(tl_amount)


def to_iso_str(value: QDate | str | None) -> str | None:
    """Convert various date inputs to ISO string."""

    if value is None:
        return None
    if isinstance(value, QDate):
        if value.isValid():
            return value.toString("yyyy-MM-dd")
        return None
    if isinstance(value, str):
        cleaned = value.strip()
        if not cleaned:
            return None
        qdate = QDate.fromString(cleaned, "dd.MM.yyyy")
        if qdate and qdate.isValid():
            return qdate.toString("yyyy-MM-dd")
        return None
    return None


def from_iso_str(value: str | None) -> QDate | None:
    """Convert ISO string to QDate."""

    if not value:
        return None
    qdate = QDate.fromString(str(value), "yyyy-MM-dd")
    if qdate and qdate.isValid():
        return qdate
    return None


def to_iso_qdate(value: QDate | None) -> str | None:
    """Convert a :class:`QDate` to ISO formatted string safely."""

    if not isinstance(value, QDate):
        return None
    if not value.isValid():
        return None
    return value.toString("yyyy-MM-dd")


def from_iso_to_qdate(value: str | None) -> QDate | None:
    """Convert an ISO formatted string into :class:`QDate`."""

    if not value:
        return None
    qdate = QDate.fromString(str(value), "yyyy-MM-dd")
    if qdate and qdate.isValid():
        return qdate
    return None


def _normalize_installment_status(
    taksit: Dict[str, Any], today: date | None = None
) -> str:
    """Ensure installment status reflects whether it is overdue.

    Paid installments keep their explicit status. For other rows the status switches
    to :data:`INSTALLMENT_OVERDUE_STATUS` when the due date is in the past.
    """

    today = today or date.today()
    status = (taksit.get("durum") or "Ödenecek").strip()
    lowered = status.casefold()
    if lowered == "ödendi":
        taksit["durum"] = status
        return status

    due_str = taksit.get("vade_tarihi")
    if due_str:
        try:
            due_date = date.fromisoformat(due_str)
        except (TypeError, ValueError):
            taksit["durum"] = status or "Ödenecek"
            return taksit["durum"]
        if due_date < today and lowered != INSTALLMENT_OVERDUE_STATUS.casefold():
            status = INSTALLMENT_OVERDUE_STATUS

    taksit["durum"] = status or "Ödenecek"
    return taksit["durum"]


def _apply_overdue_statuses(
    taksitler: List[Dict[str, Any]],
    *,
    today: date | None = None,
    cursor: sqlite3.Cursor | None = None,
) -> bool:
    """Update overdue statuses in-memory and optionally persist them.

    Returns ``True`` when at least one database row is updated.
    """

    today = today or date.today()
    any_persisted = False
    for taksit in taksitler:
        previous = taksit.get("durum") or "Ödenecek"
        updated = _normalize_installment_status(taksit, today=today)
        if cursor and updated != previous and taksit.get("id"):
            cursor.execute(
                "UPDATE taksitler SET durum = ? WHERE id = ?",
                (updated, int(taksit["id"])),
            )
            any_persisted = True
    return any_persisted


ADMIN_FORCED_PERMISSIONS: Set[str] = {"can_hard_delete"}


REQUIRED_DB_TABLES: Set[str] = {
    "dosyalar",
    "statuses",
    "attachments",
    "finans",
    "tebligatlar",
    "arabuluculuk",
}


def _calculate_column_weights(prepared_rows: List[List[Any]]) -> List[float]:
    """Tablo sütunları için göreli ağırlıkları hesaplar."""

    weights: List[float] = []
    for index, header in enumerate(HEADER_LABELS):
        max_len = len(str(header))
        for row in prepared_rows:
            value = row[index] if index < len(row) else ""
            text = "" if value is None else str(value)
            max_len = max(max_len, len(text))
        # Çok küçük değerlerin sütunu kaybetmemesi için taban ağırlık uygula
        weights.append(float(max(max_len, 3)))
    return weights


def _scale_widths(weights: List[float], available_width: float) -> List[float]:
    """Verilen ağırlıkları toplam genişliğe ölçeklendirir."""

    if not weights:
        return []
    total_weight = sum(weights)
    if total_weight <= 0:
        equal_width = available_width / len(weights)
        return [equal_width] * len(weights)
    return [available_width * (weight / total_weight) for weight in weights]


def _row_value(row: Dict[str, Any] | sqlite3.Row, key: str) -> Any:
    if isinstance(row, dict):
        return row.get(key)
    try:
        return row[key]
    except (KeyError, IndexError, TypeError):
        return None


def calculate_finance_total(finans_row: Dict[str, Any] | sqlite3.Row | None) -> int:
    """Finans kaydındaki toplam ücreti kuruş cinsinden hesaplar."""

    if not finans_row:
        return 0

    fixed_raw = _row_value(finans_row, "sozlesme_ucreti")
    fixed_decimal = Decimal("0")
    if fixed_raw not in (None, ""):
        fixed_decimal = Decimal(str(fixed_raw))
    else:
        cents_raw = _row_value(finans_row, "sozlesme_ucreti_cents")
        if cents_raw not in (None, ""):
            fixed_decimal = Decimal(str(cents_raw)) / Decimal("100")

    percent_raw = _row_value(finans_row, "sozlesme_yuzdesi")
    try:
        percent_decimal = Decimal(str(percent_raw))
    except Exception:
        percent_decimal = Decimal("0")

    target_cents = _row_value(finans_row, "tahsil_hedef_cents")
    try:
        target_decimal = Decimal(str(target_cents or 0)) / Decimal("100")
    except Exception:
        target_decimal = Decimal("0")

    total_decimal = fixed_decimal
    if percent_decimal and target_decimal:
        total_decimal += (target_decimal * percent_decimal) / Decimal("100")

    return int(
        (total_decimal * Decimal("100")).quantize(
            Decimal("1"), rounding=ROUND_HALF_UP
        )
    )


def calculate_finance_balance(finans_row: Dict[str, Any] | sqlite3.Row | None) -> int:
    """Finans kaydındaki kalan bakiyeyi kuruş cinsinden hesaplar."""

    if not finans_row:
        return 0
    total = calculate_finance_total(finans_row)
    tahsil = int(_row_value(finans_row, "tahsil_edilen_cents") or 0)
    masraf_toplam = int(_row_value(finans_row, "masraf_toplam_cents") or 0)
    masraf_tahsil = int(_row_value(finans_row, "masraf_tahsil_cents") or 0)
    return (total - tahsil) + (masraf_toplam - masraf_tahsil)


def calculate_harici_total(record: Dict[str, Any] | sqlite3.Row | None) -> int:
    """Harici finans sözleşmesinin toplam tutarını kuruş cinsinden hesaplar."""

    if not record:
        return 0
    fixed_cents = int(_row_value(record, "sabit_ucret_cents") or 0)
    target_cents = int(_row_value(record, "tahsil_hedef_cents") or 0)
    try:
        percent_rate = Decimal(str(_row_value(record, "yuzde_orani") or 0))
    except Exception:
        percent_rate = Decimal("0")
    percent_cents = 0
    # Yüzde tutarı her zaman toplama dahil (yuzde_is_sonu sadece taksitleri etkiler)
    if percent_rate and target_cents:
        percent_cents = int(
            (Decimal(target_cents) * percent_rate / Decimal("100")).quantize(
                Decimal("1"), rounding=ROUND_HALF_UP
            )
        )
    return max(fixed_cents + percent_cents, 0)


def calculate_harici_balance(record: Dict[str, Any] | sqlite3.Row | None) -> int:
    """Harici finans kaydı için kalan bakiyeyi kuruş cinsinden verir."""

    if not record:
        return 0
    total = calculate_harici_total(record)
    tahsil = int(_row_value(record, "tahsil_edilen_cents") or 0)
    masraf_toplam = int(_row_value(record, "masraf_toplam_cents") or 0)
    masraf_tahsil = int(_row_value(record, "masraf_tahsil_cents") or 0)
    balance = (total - tahsil) + (masraf_toplam - masraf_tahsil)
    return max(balance, 0)

def add_dosya(data: Dict[str, Any]) -> int:
    for date_field in ("durusma_tarihi", "is_tarihi", "is_tarihi_2"):
        if data.get(date_field) in ("", None):
            data[date_field] = None
    data.setdefault("is_archived", 0)
    conn = get_connection()
    sql = ""
    params: tuple[Any, ...] | None = None
    try:
        cur = conn.cursor()
        columns = ", ".join(data.keys())
        placeholders = ", ".join(["?"] * len(data))
        values = list(data.values())
        cur.execute(
            f"INSERT INTO dosyalar ({columns}) VALUES ({placeholders})",
            values,
        )
        dosya_id = cur.lastrowid
        cur.execute(
            "INSERT OR IGNORE INTO finans (dosya_id) VALUES (?)",
            (dosya_id,),
        )
        conn.commit()
        return dosya_id
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError(
            "Dosya kaydedilirken veritabanı hatası oluştu. Lütfen girdiğiniz bilgileri kontrol edin."
        ) from exc
    finally:
        conn.close()

def list_dosyalar() -> List[sqlite3.Row]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM dosyalar")
    rows = cur.fetchall()
    conn.close()
    return rows


def get_next_buro_takip_no() -> int:
    """Bir sonraki uygun büro takip numarasını döndürür."""
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT COALESCE(MAX(buro_takip_no), 0) + 1 FROM dosyalar")
    value = cur.fetchone()[0]
    conn.close()
    return int(value)


def get_dosya(dosya_id: int) -> Optional[Dict[str, Any]]:
    """Belirtilen kimliğe sahip tek bir dosya kaydını döndürür."""
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM dosyalar WHERE id = ?", (dosya_id,))
    row = cur.fetchone()
    conn.close()
    return dict(row) if row else None


def update_dosya(dosya_id: int, data: Dict[str, Any]) -> None:
    """Belirtilen kaydı günceller."""
    if not data:
        return
    # DÜZELTME: Sadece MEVCUT anahtarları normalize et, yeni anahtar EKLEME!
    for date_field in ("durusma_tarihi", "is_tarihi", "is_tarihi_2"):
        if date_field in data and data[date_field] in ("", None):
            data[date_field] = None
    conn = get_connection()
    try:
        cur = conn.cursor()
        columns = ", ".join(f"{key} = ?" for key in data.keys())
        values = list(data.values()) + [dosya_id]
        cur.execute(f"UPDATE dosyalar SET {columns} WHERE id = ?", values)
        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError(
            "Dosya güncellenirken veritabanı hatası oluştu. Lütfen girdileri kontrol edin."
        ) from exc
    finally:
        conn.close()


def delete_case_hard(conn: sqlite3.Connection, dosya_id: int) -> None:
    """İlişkili tüm kayıtlarla birlikte dosyayı kalıcı olarak siler."""

    import sqlite3 as _sqlite3

    def table_exists(name: str) -> bool:
        row = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
            (name,),
        ).fetchone()
        return row is not None

    try:
        conn.execute("BEGIN")

        finans_id = None
        if table_exists("finans"):
            row = conn.execute(
                "SELECT id FROM finans WHERE dosya_id=?",
                (dosya_id,),
            ).fetchone()
            finans_id = row[0] if row else None

        if finans_id is not None:
            if table_exists("taksitler"):
                conn.execute("DELETE FROM taksitler WHERE finans_id=?", (finans_id,))
            if table_exists("odeme_kayitlari"):
                conn.execute(
                    "DELETE FROM odeme_kayitlari WHERE finans_id=?",
                    (finans_id,),
                )
            if table_exists("masraflar"):
                conn.execute("DELETE FROM masraflar WHERE finans_id=?", (finans_id,))
            if table_exists("odeme_plani"):
                conn.execute("DELETE FROM odeme_plani WHERE finans_id=?", (finans_id,))
            conn.execute("DELETE FROM finans WHERE id=?", (finans_id,))
        elif table_exists("finans"):
            conn.execute("DELETE FROM finans WHERE dosya_id=?", (dosya_id,))

        if table_exists("attachments"):
            conn.execute("DELETE FROM attachments WHERE dosya_id=?", (dosya_id,))
        if table_exists("dosya_kullanicilari"):
            conn.execute(
                "DELETE FROM dosya_kullanicilari WHERE dosya_id=?",
                (dosya_id,),
            )

        conn.execute("DELETE FROM dosyalar WHERE id=?", (dosya_id,))

        conn.commit()
        return True
    except _sqlite3.Error:
        conn.rollback()
        raise

def _normalize_iso_date(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    if len(text) == 10 and text[4] == "-" and text[7] == "-":
        return text
    converted = tr_to_iso(text)
    return converted if converted else None


def _normalize_hhmm(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    parts = text.split(":")
    if len(parts) != 2:
        try:
            parsed = datetime.strptime(text, "%H%M")
        except ValueError:
            return None
        return f"{parsed.hour:02d}:{parsed.minute:02d}"
    try:
        hour, minute = (int(part) for part in parts)
    except ValueError:
        return None
    if not (0 <= hour < 24 and 0 <= minute < 60):
        return None
    return f"{hour:02d}:{minute:02d}"


def _clean_tebligat_text(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    return text or None


def get_tebligatlar_list() -> List[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, dosya_no, kurum, geldigi_tarih, teblig_tarihi, is_son_gunu, icerik,
               COALESCE(tamamlandi, 0) as tamamlandi
        FROM tebligatlar
        ORDER BY id ASC
        """
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def get_tebligat_by_id(tebligat_id: int) -> Optional[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, dosya_no, kurum, geldigi_tarih, teblig_tarihi, is_son_gunu, icerik,
               COALESCE(tamamlandi, 0) as tamamlandi
        FROM tebligatlar
        WHERE id = ?
        """,
        (tebligat_id,),
    )
    row = cur.fetchone()
    conn.close()
    return dict(row) if row else None


def insert_tebligat(rec: Dict[str, Any]) -> int:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO tebligatlar (
            dosya_no,
            kurum,
            geldigi_tarih,
            teblig_tarihi,
            is_son_gunu,
            icerik,
            updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        """,
        (
            _clean_tebligat_text(rec.get("dosya_no")),
            _clean_tebligat_text(rec.get("kurum")),
            _normalize_iso_date(rec.get("geldigi_tarih")),
            _normalize_iso_date(rec.get("teblig_tarihi")),
            _normalize_iso_date(rec.get("is_son_gunu")),
            _clean_tebligat_text(rec.get("icerik")),
        ),
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return int(new_id)


def update_tebligat(rec: Dict[str, Any]) -> None:
    tebligat_id = rec.get("id")
    if not tebligat_id:
        raise ValueError("Güncellenecek tebligat kimliği belirtilmedi")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        """
        UPDATE tebligatlar
           SET dosya_no = ?,
               kurum = ?,
               geldigi_tarih = ?,
               teblig_tarihi = ?,
               is_son_gunu = ?,
               icerik = ?,
               updated_at = CURRENT_TIMESTAMP
         WHERE id = ?
        """,
        (
            _clean_tebligat_text(rec.get("dosya_no")),
            _clean_tebligat_text(rec.get("kurum")),
            _normalize_iso_date(rec.get("geldigi_tarih")),
            _normalize_iso_date(rec.get("teblig_tarihi")),
            _normalize_iso_date(rec.get("is_son_gunu")),
            _clean_tebligat_text(rec.get("icerik")),
            int(tebligat_id),
        ),
    )
    conn.commit()
    conn.close()


def delete_tebligat(tebligat_id: int) -> None:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM tebligatlar WHERE id = ?", (tebligat_id,))
    conn.commit()
    conn.close()


def get_arabuluculuk_list() -> List[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, davaci, davali, arb_adi, arb_tel, toplanti_tarihi, toplanti_saati, konu,
               COALESCE(tamamlandi, 0) as tamamlandi
        FROM arabuluculuk
        ORDER BY id DESC
        """
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def get_arabuluculuk_by_id(rec_id: int) -> Optional[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT id, davaci, davali, arb_adi, arb_tel, toplanti_tarihi, toplanti_saati, konu,
               COALESCE(tamamlandi, 0) as tamamlandi
        FROM arabuluculuk
        WHERE id = ?
        """,
        (rec_id,),
    )
    row = cur.fetchone()
    conn.close()
    return dict(row) if row else None


def insert_arabuluculuk(rec: Dict[str, Any]) -> int:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO arabuluculuk (
            davaci,
            davali,
            arb_adi,
            arb_tel,
            toplanti_tarihi,
            toplanti_saati,
            konu,
            updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
        """,
        (
            (rec.get("davaci") or "").strip(),
            (rec.get("davali") or "").strip(),
            (rec.get("arb_adi") or "").strip(),
            (rec.get("arb_tel") or "").strip(),
            _normalize_iso_date(rec.get("toplanti_tarihi")),
            _normalize_hhmm(rec.get("toplanti_saati")) or "00:00",
            (rec.get("konu") or "").strip(),
        ),
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return int(new_id)


def update_arabuluculuk(rec: Dict[str, Any]) -> None:
    rec_id = rec.get("id")
    if not rec_id:
        raise ValueError("Güncellenecek arabuluculuk kaydı bulunamadı")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        """
        UPDATE arabuluculuk
           SET davaci = ?,
               davali = ?,
               arb_adi = ?,
               arb_tel = ?,
               toplanti_tarihi = ?,
               toplanti_saati = ?,
               konu = ?,
               updated_at = CURRENT_TIMESTAMP
         WHERE id = ?
        """,
        (
            (rec.get("davaci") or "").strip(),
            (rec.get("davali") or "").strip(),
            (rec.get("arb_adi") or "").strip(),
            (rec.get("arb_tel") or "").strip(),
            _normalize_iso_date(rec.get("toplanti_tarihi")),
            _normalize_hhmm(rec.get("toplanti_saati")) or "00:00",
            (rec.get("konu") or "").strip(),
            int(rec_id),
        ),
    )
    conn.commit()
    conn.close()


def delete_arabuluculuk(rec_id: int) -> None:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM arabuluculuk WHERE id = ?", (rec_id,))
    conn.commit()
    conn.close()


def get_attachments(dosya_id: int) -> List[Dict[str, Any]]:
    """Belirtilen dosyaya ait eklerin listesini döndürür."""

    try:
        return list_case_attachments(dosya_id)
    except Exception as exc:  # pragma: no cover - güvenlik
        logger.exception("Attachments could not be listed for case %s", dosya_id)
        raise


def add_attachment(dosya_id: int, path: str) -> int:
    """Belirtilen dosyaya yeni bir ek kaydı ekler ve dosyayı merkez klasöre kopyalar."""

    ids = store_case_attachments(dosya_id, [path])
    return ids[0] if ids else 0


def delete_attachment(attachment_id: int) -> None:
    """Ek kaydını siler."""
    remove_case_attachment(attachment_id)


def delete_attachment_with_file(attachment_id: int) -> None:
    """Ek kaydını ve ilişkili dosyayı diskte siler."""

    remove_case_attachment(attachment_id, remove_file=True)


def update_attachment_path(attachment_id: int, source_path: str) -> None:
    """Eksik dosyalar için yeni kaynak yolunu kaydeder."""

    update_case_attachment_source(attachment_id, source_path)


def export_attachment(attachment_id: int, target_directory: str) -> Path:
    """Belirtilen eki hedef klasöre kopyalar."""

    return export_case_attachment(attachment_id, target_directory)


def get_finans_master_list_bound_only(
    conn: sqlite3.Connection | None = None,
    *,
    include_archived: bool = False,
) -> List[sqlite3.Row]:
    """Return only finance records linked to dossier entries."""

    owns_conn = False
    if conn is None:
        conn = get_connection()
        owns_conn = True
    try:
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute(
            """
            WITH next_due AS (
                SELECT finans_id, MIN(vade_tarihi) AS next_due_date
                FROM taksitler
                WHERE durum != 'Ödendi'
                GROUP BY finans_id
            ),
            overdue AS (
                SELECT finans_id
                FROM taksitler
                WHERE durum != 'Ödendi' AND vade_tarihi < DATE('now')
                GROUP BY finans_id
            ),
            assignments AS (
                SELECT dosya_id, GROUP_CONCAT(user_id) AS user_ids
                FROM dosya_atamalar
                GROUP BY dosya_id
            )
            SELECT
                f.id AS finans_id,
                f.dosya_id,
                f.sozlesme_ucreti,
                f.sozlesme_ucreti_cents,
                f.sozlesme_yuzdesi,
                f.tahsil_hedef_cents,
                f.tahsil_edilen_cents,
                f.masraf_toplam_cents,
                f.masraf_tahsil_cents,
                f.yuzde_is_sonu,
                d.buro_takip_no,
                d.dosya_esas_no,
                d.muvekkil_adi,
                d.is_archived,
                next_due.next_due_date,
                CASE WHEN overdue.finans_id IS NOT NULL THEN 1 ELSE 0 END AS has_overdue_installment,
                assignments.user_ids AS assigned_user_ids
            FROM finans f
            JOIN dosyalar d ON d.id = f.dosya_id
            LEFT JOIN next_due ON next_due.finans_id = f.id
            LEFT JOIN overdue ON overdue.finans_id = f.id
            LEFT JOIN assignments ON assignments.dosya_id = f.dosya_id
            WHERE (? = 1 OR COALESCE(d.is_archived, 0) = 0)
            ORDER BY f.id DESC
            """,
            (1 if include_archived else 0,),
        )
        return cur.fetchall()
    finally:
        if owns_conn:
            conn.close()


def list_custom_tabs(conn: sqlite3.Connection) -> List[Dict[str, Any]]:
    """Özel sekmeleri isimlerine göre sıralanmış olarak döndürür."""

    cur = conn.cursor()
    cur.execute(
        "SELECT id, name, created_at FROM custom_tabs ORDER BY name COLLATE NOCASE"
    )
    rows = cur.fetchall()
    return [
        {
            "id": int(row[0]),
            "name": row[1],
            "created_at": row[2],
        }
        for row in rows
    ]


def create_custom_tab(conn: sqlite3.Connection, name: str) -> int:
    """Yeni bir özel sekme oluşturup kimliğini döndürür."""

    cleaned = name.strip()
    if not cleaned:
        raise ValueError("Sekme adı boş olamaz")

    cur = conn.cursor()
    cur.execute("INSERT INTO custom_tabs (name) VALUES (?)", (cleaned,))
    conn.commit()
    return int(cur.lastrowid)


def rename_custom_tab(conn: sqlite3.Connection, tab_id: int, new_name: str) -> None:
    """Var olan özel sekmenin adını günceller."""

    cleaned = new_name.strip()
    if not cleaned:
        raise ValueError("Sekme adı boş olamaz")

    cur = conn.cursor()
    cur.execute(
        "UPDATE custom_tabs SET name = ? WHERE id = ?",
        (cleaned, tab_id),
    )
    conn.commit()


def delete_custom_tab(conn: sqlite3.Connection, tab_id: int) -> None:
    """Özel sekmeyi kalıcı olarak siler."""

    cur = conn.cursor()
    cur.execute("DELETE FROM custom_tabs WHERE id = ?", (tab_id,))
    conn.commit()


def get_dosya_ids_for_tab(conn: sqlite3.Connection, tab_id: int) -> Set[int]:
    """Verilen sekmeye atanmış dosya kimliklerini döndürür."""

    cur = conn.cursor()
    cur.execute(
        "SELECT dosya_id FROM custom_tabs_dosyalar WHERE custom_tab_id = ?",
        (tab_id,),
    )
    return {int(row[0]) for row in cur.fetchall()}


def get_tab_assignments_for_dosya(conn: sqlite3.Connection, dosya_id: int) -> Set[int]:
    """Bir dosyanın bağlı olduğu özel sekmeleri döndürür."""

    cur = conn.cursor()
    cur.execute(
        "SELECT custom_tab_id FROM custom_tabs_dosyalar WHERE dosya_id = ?",
        (dosya_id,),
    )
    return {int(row[0]) for row in cur.fetchall()}


def set_tab_assignments_for_dosya(
    conn: sqlite3.Connection, dosya_id: int, tab_ids: Iterable[int]
) -> None:
    """Bir dosyanın özel sekme atamalarını verilen liste ile eşitler."""

    desired_ids = {int(tab_id) for tab_id in tab_ids}
    cur = conn.cursor()
    current_ids = get_tab_assignments_for_dosya(conn, dosya_id)

    to_remove = current_ids - desired_ids
    to_add = desired_ids - current_ids

    if to_remove:
        cur.executemany(
            "DELETE FROM custom_tabs_dosyalar WHERE custom_tab_id = ? AND dosya_id = ?",
            ((tab_id, dosya_id) for tab_id in to_remove),
        )

    if to_add:
        cur.executemany(
            "INSERT OR IGNORE INTO custom_tabs_dosyalar (custom_tab_id, dosya_id) VALUES (?, ?)",
            ((tab_id, dosya_id) for tab_id in to_add),
        )

    conn.commit()


def get_dosya_assignees(dosya_id: int) -> List[Dict[str, Any]]:
    """Belirtilen dosyaya atanmış kullanıcıları döndürür."""

    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        """
        SELECT u.id, u.username, u.role, u.active
        FROM dosya_atamalar da
        JOIN users u ON u.id = da.user_id
        WHERE da.dosya_id = ?
        ORDER BY LOWER(u.username)
        """,
        (dosya_id,),
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def set_dosya_assignees(dosya_id: int, user_ids: Iterable[int]) -> None:
    """Dosyaya atanmış kullanıcı listesini günceller."""

    desired: Set[int] = {int(uid) for uid in user_ids if uid is not None}
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT user_id FROM dosya_atamalar WHERE dosya_id = ?",
            (dosya_id,),
        )
        existing = {row[0] for row in cur.fetchall()}

        to_remove = existing - desired
        to_add = desired - existing

        if to_remove:
            cur.executemany(
                "DELETE FROM dosya_atamalar WHERE dosya_id = ? AND user_id = ?",
                [(dosya_id, uid) for uid in to_remove],
            )

        if to_add:
            cur.executemany(
                "INSERT OR IGNORE INTO dosya_atamalar (dosya_id, user_id) VALUES (?, ?)",
                [(dosya_id, uid) for uid in to_add],
            )

        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError(
            "Atama bilgileri güncellenirken veritabanı hatası oluştu."
        ) from exc
    finally:
        conn.close()


def set_archive_status(dosya_id: int, archived: bool) -> None:
    """Dosyanın arşiv durumunu günceller."""
    update_dosya(
        dosya_id,
        {
            "is_archived": 1 if archived else 0,
            "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        },
    )


def fetch_dosyalar_by_color_hex(
    hex6: str | None,
    search_text: str | None = None,
    open_only: bool | None = None,
    other_filters: Optional[Dict[str, Any]] = None,
    archived: bool = False,
    assigned_user_id: int | None = None,
) -> List[Dict[str, Any]]:
    """Renk koduna, arama parametrelerine ve arşiv durumuna göre kayıtları döndürür."""

    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    cur.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name='dosya_durumlari'"
    )
    has_history = cur.fetchone() is not None

    params: Dict[str, Any] = {
        "is_archived": 1 if archived else 0,
    }
    where_clauses: List[str] = ["d.is_archived = :is_archived"]

    if assigned_user_id is not None:
        params["assigned_user_id"] = int(assigned_user_id)
        where_clauses.append(
            "EXISTS (SELECT 1 FROM dosya_atamalar da WHERE da.dosya_id = d.id AND da.user_id = :assigned_user_id)"
        )

    if hex6:
        params["hex6"] = hex6
        where_clauses.append(
            "(UPPER(REPLACE(sd1.color_hex,'#','')) = :hex6 "
            "OR UPPER(REPLACE(sd2.color_hex,'#','')) = :hex6 "
            "OR UPPER(REPLACE(sa.color_hex,'#','')) = :hex6)"
        )

    durusma_period = None
    if other_filters:
        durusma_period = other_filters.get("durusma_period")

    if open_only and not archived:
        where_clauses.append(
            "NOT ("
            "UPPER(COALESCE(NULLIF(d.dava_durumu,''),'')) = 'DOSYA KAPANDI' "
            "OR UPPER(COALESCE(NULLIF(d.tekrar_dava_durumu_2,''),'')) = 'DOSYA KAPANDI'"
            ")"
        )

    if durusma_period:
        today = date.today()
        start: Optional[date] = None
        end: Optional[date] = None
        if durusma_period == "bu_hafta":
            start = today - timedelta(days=today.weekday())
            end = start + timedelta(days=6)
        elif durusma_period == "gelecek_hafta":
            start = today + timedelta(days=7 - today.weekday())
            end = start + timedelta(days=6)
        elif durusma_period == "bu_ay":
            start = today.replace(day=1)
            last_day = calendar.monthrange(today.year, today.month)[1]
            end = today.replace(day=last_day)
        if start and end:
            where_clauses.append("d.durusma_tarihi BETWEEN :durusma_start AND :durusma_end")
            params["durusma_start"] = start.isoformat()
            params["durusma_end"] = end.isoformat()

    if has_history:
        base_query = """
        WITH last_status AS (
            SELECT dd.dosya_id,
                   dd.durum AS aktif_durum,
                   MAX(dd.created_at) AS max_ts
            FROM dosya_durumlari dd
            GROUP BY dd.dosya_id
        )
        SELECT d.*,
               ls.aktif_durum AS aktif_durum,
               sa.color_hex AS status_color,
               sd1.owner AS dava_durumu_owner,
               sd2.owner AS tekrar_dava_durumu_2_owner,
               sd1.color_hex AS dava_durumu_color,
               sd2.color_hex AS tekrar_dava_durumu_2_color
        FROM dosyalar d
        LEFT JOIN last_status ls ON ls.dosya_id = d.id
        LEFT JOIN statuses sa ON TRIM(sa.ad) = TRIM(ls.aktif_durum)
        LEFT JOIN statuses sd1 ON TRIM(sd1.ad) = TRIM(d.dava_durumu)
        LEFT JOIN statuses sd2 ON TRIM(sd2.ad) = TRIM(d.tekrar_dava_durumu_2)
        """
    else:
        base_query = """
        SELECT d.*,
               COALESCE(NULLIF(d.tekrar_dava_durumu_2,''), d.dava_durumu) AS aktif_durum,
               sa.color_hex AS status_color,
               sd1.owner AS dava_durumu_owner,
               sd2.owner AS tekrar_dava_durumu_2_owner,
               sd1.color_hex AS dava_durumu_color,
               sd2.color_hex AS tekrar_dava_durumu_2_color
        FROM dosyalar d
        LEFT JOIN statuses sd1 ON TRIM(sd1.ad) = TRIM(d.dava_durumu)
        LEFT JOIN statuses sd2 ON TRIM(sd2.ad) = TRIM(d.tekrar_dava_durumu_2)
        LEFT JOIN statuses sa
          ON TRIM(sa.ad) = TRIM(COALESCE(NULLIF(d.tekrar_dava_durumu_2,''), d.dava_durumu))
        """

    if where_clauses:
        base_query += " WHERE " + " AND ".join(where_clauses)

    base_query += " ORDER BY d.buro_takip_no"

    rows = timed_query(conn, base_query, params)
    rows_prepared: list[sqlite3.Row] = list(rows)
    rows = []
    normalized_search = normalize_str(search_text) if search_text else None
    for row in rows_prepared:
        record = dict(row)
        record["status_color"] = normalize_hex(record.get("status_color"))
        record["dava_durumu_color"] = normalize_hex(record.get("dava_durumu_color"))
        record["tekrar_dava_durumu_2_color"] = normalize_hex(
            record.get("tekrar_dava_durumu_2_color")
        )
        record["dava_durumu_owner"] = resolve_owner_label(
            record.get("dava_durumu_owner"), record.get("dava_durumu_color")
        )
        record["tekrar_dava_durumu_2_owner"] = resolve_owner_label(
            record.get("tekrar_dava_durumu_2_owner"),
            record.get("tekrar_dava_durumu_2_color"),
        )
        if normalized_search:
            hay_fields = [
                "dosya_esas_no",
                "muvekkil_adi",
                "karsi_taraf",
                "dosya_konusu",
                "mahkeme_adi",
                "dava_durumu",
                "aciklama",
                "tekrar_dava_durumu_2",
                "aciklama_2",
            ]
            haystack = " ".join(
                normalize_str(str(record.get(field, ""))) for field in hay_fields
            )
            if normalized_search not in haystack:
                continue
        rows.append(record)

    conn.close()
    return rows


def get_all_dosyalar(
    archived: bool = False, assigned_user_id: int | None = None
) -> List[Dict[str, Any]]:
    """`dosyalar` tablosundaki tüm kayıtları döndürür."""

    return fetch_dosyalar_by_color_hex(
        None,
        archived=archived,
        assigned_user_id=assigned_user_id,
    )


def search_dosyalar(filters: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Verilen filtre sözlüğüne göre dosya kayıtlarını döndürür."""
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    base_query = (
        "SELECT id, buro_takip_no, dosya_esas_no, muvekkil_adi, muvekkil_rolu, "
        "karsi_taraf, dosya_konusu, mahkeme_adi, dava_acilis_tarihi, "
        "durusma_tarihi, dava_durumu, is_tarihi, aciklama, tekrar_dava_durumu_2, "
        "is_tarihi_2, aciklama_2 FROM dosyalar"
    )

    conditions: List[str] = ["is_archived = 0"]
    params: List[Any] = []

    q = filters.get("query")
    if q:
        like = f"%{q.lower()}%"
        text_fields = [
            "dosya_esas_no",
            "muvekkil_adi",
            "karsi_taraf",
            "dosya_konusu",
            "mahkeme_adi",
            "dava_durumu",
            "aciklama",
            "tekrar_dava_durumu_2",
            "aciklama_2",
        ]
        conditions.append(
            "(" + " OR ".join(f"LOWER({field}) LIKE ?" for field in text_fields) + ")"
        )
        params.extend([like] * len(text_fields))

    rol = filters.get("rol")
    if rol:
        conditions.append("muvekkil_rolu = ?")
        params.append(rol)

    status = filters.get("status")
    if status:
        conditions.append("dava_durumu = ?")
        params.append(status)

    period = filters.get("durusma_period")
    if period:
        today = date.today()
        if period == "bu_hafta":
            start = today - timedelta(days=today.weekday())
            end = start + timedelta(days=6)
        elif period == "gelecek_hafta":
            start = today + timedelta(days=7 - today.weekday())
            end = start + timedelta(days=6)
        elif period == "bu_ay":
            start = today.replace(day=1)
            last_day = calendar.monthrange(today.year, today.month)[1]
            end = today.replace(day=last_day)
        else:
            start = end = None
        if start and end:
            conditions.append("durusma_tarihi BETWEEN ? AND ?")
            params.extend([start.isoformat(), end.isoformat()])

    if filters.get("only_open"):
        conditions.append("dava_durumu <> 'Kapandı'")

    if conditions:
        base_query += " WHERE " + " AND ".join(conditions)

    base_query += " ORDER BY buro_takip_no"

    cur.execute(base_query, params)
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows

def add_status(ad: str, color_hex: str, owner: str) -> int:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO statuses (ad, color_hex, owner) VALUES (?, ?, ?)",
        (ad, normalize_hex(color_hex) or color_hex, owner),
    )
    conn.commit()
    rowid = cur.lastrowid
    conn.close()
    return rowid

def get_statuses() -> List[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM statuses")
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


# Geriye dönük uyumluluk için
list_statuses = get_statuses


def update_status(status_id: int, ad: str, color_hex: str, owner: str) -> None:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        "UPDATE statuses SET ad = ?, color_hex = ?, owner = ? WHERE id = ?",
        (ad, normalize_hex(color_hex) or color_hex, owner, status_id),
    )
    conn.commit()
    conn.close()


def delete_status(status_id: int) -> None:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM statuses WHERE id = ?", (status_id,))
    conn.commit()
    conn.close()


def get_status_color(status_ad: str) -> Optional[str]:
    """Verilen statü adının hex renk kodunu döndürür.

    Türkçe karakterleri doğru şekilde karşılaştırır (İ/i, I/ı).
    """
    if not status_ad:
        return None
    conn = get_connection()
    cur = conn.cursor()
    # Tüm statüleri al ve Python'da Türkçe karakter destekli karşılaştırma yap
    cur.execute("SELECT ad, color_hex FROM statuses")
    rows = cur.fetchall()
    conn.close()

    # Aranan statü adını normalize et
    search_normalized = turkish_casefold(status_ad.strip())

    for row in rows:
        ad, color_hex = row[0], row[1]
        if ad and turkish_casefold(ad.strip()) == search_normalized:
            return normalize_hex(color_hex) if color_hex else None

    return None


def get_settings(key: str) -> Optional[str]:
    """Ayarlar tablosundan verilen anahtarın değerini döndürür."""
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT value FROM ayarlar WHERE key = ?", (key,))
    row = cur.fetchone()
    conn.close()
    return row[0] if row else None


def set_settings(key: str, value: str) -> None:
    """Ayar değerini ekler veya günceller."""
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO ayarlar (key, value) VALUES (?, ?)
        ON CONFLICT(key) DO UPDATE SET value = excluded.value
        """,
        (key, value),
    )
    conn.commit()
    conn.close()


# Tekil adlarla uyumlu yardımcılar
def get_setting(key: str) -> Optional[str]:
    return get_settings(key)


def set_setting(key: str, value: str) -> None:
    set_settings(key, value)


EXPORT_FIELDS = [
    "buro_takip_no",
    "dosya_esas_no",
    "muvekkil_adi",
    "muvekkil_rolu",
    "karsi_taraf",
    "dosya_konusu",
    "mahkeme_adi",
    "dava_acilis_tarihi",
    "durusma_tarihi",
    "dava_durumu",
    "is_tarihi",
    "aciklama",
    "tekrar_dava_durumu_2",
    "is_tarihi_2",
    "aciklama_2",
]


EXPORT_HEADERS: dict[str, str] = {
    "buro_takip_no": "Büro Takip No",
    "dosya_esas_no": "Dosya Esas No",
    "muvekkil_adi": "Müvekkil Adı",
    "muvekkil_rolu": "Müvekkil Rolü",
    "karsi_taraf": "Karşı Taraf",
    "dosya_konusu": "Dosya Konusu",
    "mahkeme_adi": "Mahkeme Adı",
    "dava_acilis_tarihi": "Dava Açılış Tarihi",
    "durusma_tarihi": "Duruşma Tarihi",
    "dava_durumu": "Dava Durumu",
    "is_tarihi": "İş Tarihi",
    "aciklama": "Açıklama",
    "tekrar_dava_durumu_2": "Dava Durumu 2",
    "is_tarihi_2": "İş Tarihi 2",
    "aciklama_2": "Açıklama 2",
}


DATE_FIELDS = {"dava_acilis_tarihi", "durusma_tarihi", "is_tarihi", "is_tarihi_2"}


HEADER_LABELS = [EXPORT_HEADERS.get(field, field) for field in EXPORT_FIELDS]


STATUS_COLOR_FIELDS = [
    ("dava_durumu_color", EXPORT_FIELDS.index("dava_durumu")),
    ("tekrar_dava_durumu_2_color", EXPORT_FIELDS.index("tekrar_dava_durumu_2")),
]

def _format_export_value(field: str, value: Any) -> Any:
    if value is None:
        return ""
    if field in DATE_FIELDS and value:
        return iso_to_tr(value)
    return value


def _prepare_export_row(row: Dict[str, Any]) -> List[Any]:
    return [_format_export_value(key, row.get(key)) for key in EXPORT_FIELDS]


def _apply_docx_cell_colors(cell, bg_hex: str | None) -> None:
    normalized_bg = normalize_hex(bg_hex)
    if not normalized_bg:
        return
    text_color = get_status_text_color(normalized_bg)
    normalized_text = normalize_hex(text_color) or "000000"

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), normalized_bg)
    shading.set(qn("w:color"), normalized_text)

    rgb = RGBColor.from_string(normalized_text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.color.rgb = rgb


def _prepare_export_dict(row: Dict[str, Any]) -> Dict[str, Any]:
    values = _prepare_export_row(row)
    return dict(zip(HEADER_LABELS, values))


def export_dosyalar_to_csv(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini CSV olarak dışa aktarır."""
    data = [_prepare_export_dict(row) for row in rows]
    df = pd.DataFrame(data, columns=HEADER_LABELS)
    df.to_csv(path, index=False, encoding="utf-8")


def export_dosyalar_to_xlsx(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini XLSX olarak dışa aktarır."""
    wb = Workbook()
    ws = wb.active
    ws.append(HEADER_LABELS)
    for row in rows:
        ws.append(_prepare_export_row(row))
    wb.save(path)


def export_dosyalar_to_docx(path: str, rows: List[Dict[str, Any]]) -> None:
    """Verilen kayıt listesini Word belgesi olarak dışa aktarır."""
    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(9)

    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width, section.page_height = new_width, new_height

    prepared_rows = [_prepare_export_row(row) for row in rows]
    column_weights = _calculate_column_weights(prepared_rows)

    table = doc.add_table(rows=1, cols=len(HEADER_LABELS))
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for index, header in enumerate(HEADER_LABELS):
        cell = header_cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.style = normal_style
        paragraph.text = ""
        run = paragraph.add_run(str(header))
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"

    for row, prepared_row in zip(rows, prepared_rows):
        row_cells = table.add_row().cells
        for index, value in enumerate(prepared_row):
            text = "" if value is None else str(value)
            cell = row_cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.style = normal_style
            paragraph.text = ""
            lines = text.split("\n")
            for idx, part in enumerate(lines):
                run = paragraph.add_run(part)
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                if idx < len(lines) - 1:
                    run.add_break()
        for field, column_index in STATUS_COLOR_FIELDS:
            color_hex = normalize_hex(row.get(field))
            if not color_hex:
                continue
            _apply_docx_cell_colors(row_cells[column_index], color_hex)

    if column_weights:
        first_section = doc.sections[0]
        available_width = (
            first_section.page_width
            - first_section.left_margin
            - first_section.right_margin
        )
        col_widths = _scale_widths(column_weights, float(available_width))
        for col_index, width in enumerate(col_widths):
            table.columns[col_index].width = int(width)
        for row in table.rows:
            for col_index, cell in enumerate(row.cells):
                cell.width = int(col_widths[col_index])

    doc.save(path)


def backup_database(dest_path: str) -> None:
    """Veritabanı dosyasını belirtilen yola kopyalar."""
    shutil.copy2(DB_PATH, dest_path)


def validate_database_file(
    path: str, required_tables: Iterable[str] | None = None
) -> tuple[bool, Set[str]]:
    """Seçilen veritabanı dosyasının gerekli tabloları içerip içermediğini kontrol eder."""

    required = {table.lower() for table in (required_tables or REQUIRED_DB_TABLES)}
    if not path or not os.path.exists(path):
        return False, required

    try:
        conn = sqlite3.connect(path)
        cur = conn.cursor()
        cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = {row[0].lower() for row in cur.fetchall()}
    except sqlite3.Error:
        return False, required
    finally:
        try:
            conn.close()
        except Exception:  # pragma: no cover - best effort
            pass

    missing = {table for table in required if table not in tables}
    return not missing, missing


def get_permissions_for_role(role: str) -> Dict[str, bool]:
    """Verilen rol için izin haritasını döndürür."""

    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        "SELECT action, allowed FROM permissions WHERE role = ?",
        (role,),
    )
    permissions = {action: bool(allowed) for action, allowed in cur.fetchall()}
    conn.close()

    defaults = DEFAULT_ROLE_PERMISSIONS.get(role, {})
    result: Dict[str, bool] = {}
    for action in PERMISSION_ACTIONS:
        if action in permissions:
            result[action] = permissions[action]
        else:
            result[action] = bool(defaults.get(action, False))

    if role == "admin":
        for forced_action in ADMIN_FORCED_PERMISSIONS:
            result[forced_action] = True

    return result


def get_all_permissions() -> Dict[str, Dict[str, bool]]:
    """Tüm roller için izinleri döndürür."""

    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT role, action, allowed FROM permissions")
    mapping: Dict[str, Dict[str, bool]] = {}
    for role, action, allowed in cur.fetchall():
        mapping.setdefault(role, {})[action] = bool(allowed)
    conn.close()

    for role_value, _ in USER_ROLE_CHOICES:
        role_defaults = DEFAULT_ROLE_PERMISSIONS.get(role_value, {})
        role_map = mapping.setdefault(role_value, {})
        for action in PERMISSION_ACTIONS:
            role_map.setdefault(action, bool(role_defaults.get(action, False)))
        if role_value == "admin":
            for forced_action in ADMIN_FORCED_PERMISSIONS:
                role_map[forced_action] = True
    return mapping


def set_permissions_for_role(role: str, permissions: Dict[str, bool]) -> None:
    """Belirtilen rol için izinleri günceller."""

    conn = get_connection()
    cur = conn.cursor()
    for action in PERMISSION_ACTIONS:
        if action not in permissions:
            continue
        if role == "admin" and action in ADMIN_FORCED_PERMISSIONS:
            continue
        cur.execute(
            """
            INSERT INTO permissions (role, action, allowed)
            VALUES (?, ?, ?)
            ON CONFLICT(role, action) DO UPDATE SET allowed = excluded.allowed
            """,
            (role, action, 1 if permissions[action] else 0),
        )
    conn.commit()
    conn.close()


# Kullanıcı işlemleri

def authenticate(username: str, password: str) -> Optional[Dict[str, Any]]:
    """Kullanıcı adı ve parola ile giriş doğrulaması yapar."""
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        "SELECT * FROM users WHERE username = ? AND active = 1", (username,)
    )
    row = cur.fetchone()
    conn.close()
    if row and verify_password(password, row["password_hash"]):
        user_dict = dict(row)
        role = user_dict.get("role", "") or ""
        user_dict["permissions"] = get_permissions_for_role(role)
        return user_dict
    return None


def add_user(username: str, password: str, role: str, active: bool) -> int:
    """Yeni kullanıcı ekler."""
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO users (username, password_hash, role, active) VALUES (?, ?, ?, ?)",
        (username, hash_password(password), role, 1 if active else 0),
    )
    conn.commit()
    uid = cur.lastrowid
    conn.close()
    return uid


def update_user(
    user_id: int,
    username: str,
    password: str | None = None,
    role: str | None = None,
    active: bool | None = None,
) -> None:
    """Kullanıcı bilgilerini günceller."""
    conn = get_connection()
    cur = conn.cursor()
    set_parts: List[str] = []
    params: List[Any] = []

    set_parts.append("username = ?")
    params.append(username)

    if password:
        set_parts.append("password_hash = ?")
        params.append(hash_password(password))
    if role is not None:
        set_parts.append("role = ?")
        params.append(role)
    if active is not None:
        set_parts.append("active = ?")
        params.append(1 if active else 0)

    set_parts.append("updated_at = CURRENT_TIMESTAMP")
    query = f"UPDATE users SET {', '.join(set_parts)} WHERE id = ?"
    params.append(user_id)
    cur.execute(query, params)
    conn.commit()
    conn.close()


def delete_user(user_id: int) -> None:
    """Kullanıcıyı siler. Admin kullanıcı (id=1) silinemez."""
    if user_id == 1:
        raise ValueError("Admin kullanıcı silinemez.")
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM users WHERE id = ?", (user_id,))
    conn.commit()
    conn.close()


def get_users() -> List[Dict[str, Any]]:
    """Tüm kullanıcıları listeler."""
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        "SELECT id, username, role, active, created_at, updated_at FROM users"
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def log_action(user_id: int, action: str, target_id: Optional[int] = None) -> None:
    """Audit log tablosuna bir işlem kaydı ekler."""
    conn = get_connection()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO audit_log (user_id, action, target_id) VALUES (?, ?, ?)",
        (user_id, action, target_id),
    )
    conn.commit()
    conn.close()


def ensure_finans_record(
    dosya_id: int, cur: sqlite3.Cursor | None = None
) -> int:
    """Verilen dosya için finans kaydı oluşturur veya mevcut olanı döndürür."""

    owns_connection = cur is None
    conn: sqlite3.Connection | None = None
    if owns_connection:
        conn = get_connection()
        cur = conn.cursor()
    assert cur is not None
    cur.execute(
        "INSERT OR IGNORE INTO finans (dosya_id) VALUES (?)",
        (dosya_id,),
    )
    cur.execute("SELECT id FROM finans WHERE dosya_id = ?", (dosya_id,))
    row = cur.fetchone()
    finans_id = int(row[0]) if row else 0
    if owns_connection and conn is not None:
        conn.commit()
        conn.close()
    return finans_id


def _hydrate_finans_row(
    row: sqlite3.Row | Dict[str, Any],
    *,
    fallback_dosya_id: int | None = None,
) -> Dict[str, Any]:
    record = dict(row)
    record.setdefault("dosya_id", fallback_dosya_id)
    finans_id = record.get("id") or record.get("finans_id")
    try:
        record["finans_id"] = int(finans_id) if finans_id is not None else None
    except (TypeError, ValueError):
        record["finans_id"] = None
    if record.get("finans_id") is None and fallback_dosya_id is not None:
        record["finans_id"] = record.get("id")

    record["toplam_ucret_cents"] = calculate_finance_total(row)
    record["kalan_bakiye_cents"] = calculate_finance_balance(row)
    record["yuzde_is_sonu"] = int(record.get("yuzde_is_sonu") or 0)
    fixed_value = record.get("sozlesme_ucreti")
    if fixed_value not in (None, ""):
        try:
            record["sozlesme_ucreti"] = float(fixed_value)
        except (TypeError, ValueError):
            record["sozlesme_ucreti"] = float(Decimal(str(fixed_value)))
    else:
        cents_value = record.get("sozlesme_ucreti_cents")
        if cents_value not in (None, ""):
            record["sozlesme_ucreti"] = float(
                Decimal(str(cents_value)) / Decimal("100")
            )
        else:
            record["sozlesme_ucreti"] = 0.0
    record["sozlesme_ucreti_cents"] = tl_to_cents(record.get("sozlesme_ucreti") or 0)

    percent_value = record.get("sozlesme_yuzdesi")
    if percent_value in (None, "", False):
        record["sozlesme_yuzdesi"] = 0.0
    else:
        try:
            record["sozlesme_yuzdesi"] = float(percent_value)
        except (TypeError, ValueError, InvalidOperation):
            record["sozlesme_yuzdesi"] = 0.0
    return record


def get_finans_for_dosya(dosya_id: int) -> Dict[str, Any]:
    """Belirtilen dosya için finans kaydını döndürür."""

    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    finans_id = ensure_finans_record(dosya_id, cur)
    cur.execute("SELECT * FROM finans WHERE dosya_id = ?", (dosya_id,))
    row = cur.fetchone()
    conn.commit()
    conn.close()
    if not row:
        return {
            "dosya_id": dosya_id,
            "id": finans_id,
            "finans_id": finans_id,
        }
    return _hydrate_finans_row(row, fallback_dosya_id=dosya_id)


def get_finans_by_id(finans_id: int) -> Dict[str, Any]:
    """Fetch finance information by its identifier."""

    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM finans WHERE id = ?", (finans_id,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return {"finans_id": finans_id, "id": finans_id}
    return _hydrate_finans_row(row)


def list_finance_overview(include_archived: bool = False) -> List[Dict[str, Any]]:
    """Tüm finans kayıtlarının özetini döndürür."""

    conn = get_connection()
    conn.row_factory = sqlite3.Row
    try:
        cur = conn.cursor()
        cur.execute("INSERT OR IGNORE INTO finans (dosya_id) SELECT id FROM dosyalar")
        conn.commit()
        rows = get_finans_master_list_bound_only(
            conn, include_archived=include_archived
        )
    finally:
        conn.close()

    overview: List[Dict[str, Any]] = []
    for row in rows:
        record = dict(row)
        finans_id = record.get("finans_id")
        try:
            record["finans_id"] = int(finans_id) if finans_id is not None else None
        except (TypeError, ValueError):
            record["finans_id"] = None

        dosya_id = record.get("dosya_id")
        if dosya_id is not None and dosya_id != "":
            try:
                record["dosya_id"] = int(dosya_id)
            except (TypeError, ValueError):
                record["dosya_id"] = None
        else:
            record["dosya_id"] = None

        bn_value = record.get("buro_takip_no")
        record["buro_takip_no"] = bn_value if bn_value not in (None, "") else ""
        record["dosya_esas_no"] = record.get("dosya_esas_no")
        record["muvekkil_adi"] = record.get("muvekkil_adi")

        raw_user_ids = record.get("assigned_user_ids")
        if raw_user_ids:
            record["assigned_user_ids"] = [
                int(user_id)
                for user_id in str(raw_user_ids).split(",")
                if user_id
            ]
        else:
            record["assigned_user_ids"] = []

        record["toplam_ucret_cents"] = calculate_finance_total(row)
        record["kalan_bakiye_cents"] = calculate_finance_balance(row)
        record["has_overdue_installment"] = bool(record.get("has_overdue_installment"))
        record["yuzde_is_sonu"] = int(record.get("yuzde_is_sonu") or 0)
        record["due_category"] = _categorize_due_date(record.get("next_due_date"))

        fixed_value = record.get("sozlesme_ucreti")
        if fixed_value not in (None, ""):
            try:
                record["sozlesme_ucreti"] = float(fixed_value)
            except (TypeError, ValueError):
                record["sozlesme_ucreti"] = float(Decimal(str(fixed_value)))
        else:
            cents_value = record.get("sozlesme_ucreti_cents")
            if cents_value not in (None, ""):
                record["sozlesme_ucreti"] = float(
                    Decimal(str(cents_value)) / Decimal("100")
                )
            else:
                record["sozlesme_ucreti"] = 0.0
        record["sozlesme_ucreti_cents"] = tl_to_cents(record.get("sozlesme_ucreti") or 0)

        percent_value = record.get("sozlesme_yuzdesi")
        if percent_value in (None, "", False):
            record["sozlesme_yuzdesi"] = 0.0
        else:
            try:
                record["sozlesme_yuzdesi"] = float(percent_value)
            except (TypeError, ValueError, InvalidOperation):
                record["sozlesme_yuzdesi"] = 0.0

        overview.append(record)
    return overview


def summarize_finance_by_ids(finance_ids: Iterable[int]) -> Dict[str, int]:
    """Verilen finans kayıtları için toplamları tek sorguda hesaplar."""

    normalized: list[int] = []
    for finance_id in finance_ids:
        try:
            normalized.append(int(finance_id))
        except (TypeError, ValueError):
            continue
    normalized = sorted(set(normalized))
    if not normalized:
        return {"contract": 0, "collected": 0, "expense": 0, "balance": 0}

    placeholders = ",".join(["?"] * len(normalized))
    query = f"""
        WITH base AS (
            SELECT
                f.id AS finans_id,
                (
                    COALESCE(
                        CASE
                            WHEN f.sozlesme_ucreti IS NOT NULL
                                 AND TRIM(CAST(f.sozlesme_ucreti AS TEXT)) != ''
                                THEN CAST(ROUND(f.sozlesme_ucreti * 100.0) AS INTEGER)
                            WHEN f.sozlesme_ucreti_cents IS NOT NULL
                                 AND TRIM(CAST(f.sozlesme_ucreti_cents AS TEXT)) != ''
                                THEN CAST(f.sozlesme_ucreti_cents AS INTEGER)
                            ELSE 0
                        END,
                        0
                    )
                    +
                    CASE
                        WHEN COALESCE(f.sozlesme_yuzdesi, 0) != 0
                             AND COALESCE(f.tahsil_hedef_cents, 0) != 0
                            THEN CAST(
                                ROUND((f.tahsil_hedef_cents * f.sozlesme_yuzdesi) / 100.0)
                                AS INTEGER
                            )
                        ELSE 0
                    END
                ) AS toplam_ucret_cents,
                COALESCE(f.tahsil_edilen_cents, 0) AS tahsil_edilen_cents,
                COALESCE(f.masraf_toplam_cents, 0) AS masraf_toplam_cents,
                COALESCE(f.masraf_tahsil_cents, 0) AS masraf_tahsil_cents
            FROM finans f
            WHERE f.id IN ({placeholders})
        )
        SELECT
            COALESCE(SUM(toplam_ucret_cents), 0) AS toplam_sozlesme_cents,
            COALESCE(SUM(tahsil_edilen_cents), 0) AS toplam_tahsil_cents,
            COALESCE(SUM(masraf_toplam_cents), 0) AS toplam_masraf_cents,
            COALESCE(
                SUM(
                    toplam_ucret_cents - tahsil_edilen_cents
                    + masraf_toplam_cents - masraf_tahsil_cents
                ),
                0
            ) AS kalan_cents
        FROM base
    """

    conn = get_connection()
    try:
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute(query, normalized)
        row = cur.fetchone()
    finally:
        conn.close()

    if not row:
        return {"contract": 0, "collected": 0, "expense": 0, "balance": 0}

    def _extract(value: Any) -> int:
        try:
            return int(value or 0)
        except (TypeError, ValueError):
            return 0

    contract = _extract(row["toplam_sozlesme_cents"])
    collected = _extract(row["toplam_tahsil_cents"])
    expense = _extract(row["toplam_masraf_cents"])
    balance = _extract(row["kalan_cents"])
    return {
        "contract": contract,
        "collected": collected,
        "expense": expense,
        "balance": balance,
    }


def summarize_harici_finance_by_ids(finance_ids: Iterable[int]) -> Dict[str, int]:
    """External finance tabloları için toplu özet döndür."""

    normalized: list[int] = []
    for finance_id in finance_ids:
        try:
            normalized.append(int(finance_id))
        except (TypeError, ValueError):
            continue
    normalized = sorted(set(normalized))
    if not normalized:
        return {"contract": 0, "collected": 0, "expense": 0, "balance": 0}

    placeholders = ",".join(["?"] * len(normalized))
    query = f"""
        SELECT
            COALESCE(
                SUM(
                    CASE
                        WHEN COALESCE(f.toplam_ucret_cents, 0) > 0 THEN f.toplam_ucret_cents
                        ELSE COALESCE(f.sabit_ucret_cents, 0)
                             + CASE
                                 WHEN COALESCE(f.yuzde_is_sonu, 0) = 0
                                      AND COALESCE(f.yuzde_orani, 0) != 0
                                      AND COALESCE(f.tahsil_hedef_cents, 0) != 0
                                     THEN CAST(ROUND(f.tahsil_hedef_cents * f.yuzde_orani / 100.0) AS INTEGER)
                                 ELSE 0
                               END
                    END
                ),
                0
            ) AS contract_cents,
            COALESCE(SUM(f.tahsil_edilen_cents), 0) AS collected_cents,
            COALESCE(SUM(f.masraf_toplam_cents), 0) AS expense_cents,
            COALESCE(SUM(f.masraf_tahsil_cents), 0) AS expense_collected_cents
        FROM finans_harici f
        WHERE f.id IN ({placeholders})
    """

    conn = get_connection()
    try:
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute(query, normalized)
        row = cur.fetchone()
    finally:
        conn.close()

    if not row:
        return {"contract": 0, "collected": 0, "expense": 0, "balance": 0}

    def _extract(value: Any) -> int:
        try:
            return int(value or 0)
        except (TypeError, ValueError):
            return 0

    contract = _extract(row["contract_cents"])
    collected = _extract(row["collected_cents"])
    expense = _extract(row["expense_cents"])
    expense_collected = _extract(row["expense_collected_cents"])
    balance = contract + expense - collected - expense_collected
    return {
        "contract": contract,
        "collected": collected,
        "expense": expense,
        "balance": balance,
    }


def _ensure_conn(conn: sqlite3.Connection | None) -> tuple[sqlite3.Connection, bool]:
    owns = False
    if conn is None:
        conn = get_connection()
        owns = True
    return conn, owns


def _coerce_auto_payment_date(value: Any) -> str | None:
    """Normalize various date inputs to ISO format for auto payments."""

    if value is None:
        return None
    if isinstance(value, QDate):
        if value.isValid():
            return value.toString("yyyy-MM-dd")
        return None
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, str):
        trimmed = value.strip()
        if not trimmed:
            return None
        qdate = QDate.fromString(trimmed, "dd.MM.yyyy")
        if qdate and qdate.isValid():
            return qdate.toString("yyyy-MM-dd")
        return trimmed
    return None


def harici_get_contract(conn, hid: int):
    """Fetch primary contract fields for an external finance record."""

    if conn is None:
        return None
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT
            id,
            harici_bn,
            harici_muvekkil,
            harici_esas_no,
            COALESCE(sabit_ucret_cents, 0) AS sabit_ucret_cents,
            COALESCE(yuzde_orani, 0) AS yuzde_orani,
            COALESCE(tahsil_hedef_cents, 0) AS tahsil_hedef_cents,
            COALESCE(yuzde_is_sonu, 0) AS yuzde_is_sonu,
            COALESCE(masraf_toplam_cents, 0) AS masraf_toplam_cents,
            COALESCE(masraf_tahsil_cents, 0) AS masraf_tahsil_cents,
            COALESCE(tahsil_edilen_cents, 0) AS tahsil_edilen_cents,
            COALESCE(toplam_ucret_cents, 0) AS toplam_ucret_cents,
            COALESCE(kalan_bakiye_cents, 0) AS kalan_bakiye_cents,
            plan_taksit_sayisi,
            plan_periyot,
            plan_vade_gunu,
            plan_baslangic_tarihi,
            plan_aciklama,
            notlar
        FROM finans_harici
        WHERE id = ?
        """,
        (hid,),
    )
    row = cursor.fetchone()
    if row is None:
        return None
    columns = [col[0] for col in cursor.description or []]
    result: dict[str, Any] = {}
    if isinstance(row, sqlite3.Row):
        for column in columns:
            try:
                result[column] = row[column]
            except (KeyError, IndexError, TypeError):
                continue
        return result
    for idx, column in enumerate(columns):
        if idx < len(row):
            result[column] = row[idx]
    return result


def harici_update_contract(
    conn,
    hid: int,
    *,
    sabit_ucret_cents: int,
    yuzde_orani: float,
    tahsil_hedef_cents: int,
    yuzde_is_sonu: int,
    notlar: str | None,
    harici_bn: str | None,
    harici_muvekkil: str | None,
    harici_esas_no: str | None,
    karsi_vekalet_ucreti_cents: int = 0,
) -> None:
    """Persist contract information for an external finance record."""

    if conn is None:
        raise ValueError("Veritabanı bağlantısı bulunamadı.")
    with conn:
        conn.execute(
            """
            UPDATE finans_harici
               SET sabit_ucret_cents = ?,
                   yuzde_orani = ?,
                   tahsil_hedef_cents = ?,
                   yuzde_is_sonu = ?,
                   notlar = ?,
                   harici_bn = ?,
                   harici_muvekkil = ?,
                   harici_esas_no = ?,
                   karsi_vekalet_ucreti_cents = ?,
                   updated_at = CURRENT_TIMESTAMP
             WHERE id = ?
            """,
            (
                int(sabit_ucret_cents or 0),
                float(yuzde_orani or 0.0),
                int(tahsil_hedef_cents or 0),
                1 if yuzde_is_sonu else 0,
                notlar or None,
                harici_bn or None,
                harici_muvekkil or None,
                harici_esas_no or None,
                int(karsi_vekalet_ucreti_cents or 0),
                hid,
            ),
        )


def harici_generate_installments(
    conn,
    hid: int,
    taksit_sayisi: int,
    periyot: str,
    vade_gunu: int,
    baslangic_tarihi: date | None,
) -> list[dict[str, Any]]:
    """Calculate installment rows for an external finance record."""

    if taksit_sayisi <= 0:
        return []
    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT COALESCE(sabit_ucret_cents,0) AS sabit,
                   COALESCE(yuzde_orani,0) AS oran,
                   COALESCE(tahsil_hedef_cents,0) AS hedef,
                   COALESCE(yuzde_is_sonu,0) AS ertelenmis
              FROM finans_harici
             WHERE id=?
            """,
            (hid,),
        )
        row = cursor.fetchone()
        if not row:
            return []
        if isinstance(row, sqlite3.Row):
            record = dict(row)
        else:
            record = {
                "sabit": row[0] if len(row) > 0 else 0,
                "oran": row[1] if len(row) > 1 else 0,
                "hedef": row[2] if len(row) > 2 else 0,
                "ertelenmis": row[3] if len(row) > 3 else 0,
            }
        fixed_cents = int(record.get("sabit") or 0)
        percent_rate = Decimal(str(record.get("oran") or 0))
        target_cents = int(record.get("hedef") or 0)
        deferred = int(record.get("ertelenmis") or 0)
    finally:
        if owns:
            conn.close()

    percent_cents = 0
    if percent_rate and target_cents:
        percent_cents = int(
            (Decimal(target_cents) * percent_rate / Decimal("100")).quantize(
                Decimal("1"), rounding=ROUND_HALF_UP
            )
        )
    plan_total = fixed_cents
    if not deferred:
        plan_total += percent_cents
    if plan_total <= 0:
        return []

    start_date = baslangic_tarihi or date.today()
    amounts: list[int] = []
    base = plan_total // taksit_sayisi
    remainder = plan_total - (base * taksit_sayisi)
    for index in range(taksit_sayisi):
        value = base
        if index == taksit_sayisi - 1:
            value += remainder
        amounts.append(value)

    taksitler: list[dict[str, Any]] = []
    normalized_period = (periyot or "Ay").lower()
    normalized_day = max(1, min(28, int(vade_gunu or 1)))
    for idx, amount in enumerate(amounts):
        if normalized_period == "ay":
            if idx == 0:
                due = start_date
            else:
                due = _add_months(start_date, idx, normalized_day)
        else:
            due = start_date + timedelta(days=7 * idx)
        taksitler.append(
            {
                "sira": idx + 1,
                "vade_tarihi": due.strftime("%Y-%m-%d"),
                "tutar_cents": amount,
                "durum": "Ödenecek",
                "odeme_tarihi": None,
                "aciklama": "",
            }
        )
    return taksitler


def harici_get_payment_plan(
    conn, hid: int
) -> tuple[dict[str, Any] | None, list[dict[str, Any]]]:
    """Fetch persisted plan metadata and installment rows."""

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT plan_taksit_sayisi, plan_periyot, plan_vade_gunu,
                   plan_baslangic_tarihi, plan_aciklama
              FROM finans_harici
             WHERE id=?
            """,
            (hid,),
        )
        plan_row = cursor.fetchone()
        plan = None
        if plan_row:
            if isinstance(plan_row, sqlite3.Row):
                plan = dict(plan_row)
            else:
                plan = {
                    "plan_taksit_sayisi": plan_row[0] if len(plan_row) > 0 else 0,
                    "plan_periyot": plan_row[1] if len(plan_row) > 1 else None,
                    "plan_vade_gunu": plan_row[2] if len(plan_row) > 2 else 0,
                    "plan_baslangic_tarihi": plan_row[3] if len(plan_row) > 3 else None,
                    "plan_aciklama": plan_row[4] if len(plan_row) > 4 else None,
                }
            plan = {
                "taksit_sayisi": int(plan.get("plan_taksit_sayisi") or 0),
                "periyot": plan.get("plan_periyot") or "Ay",
                "vade_gunu": int(plan.get("plan_vade_gunu") or 1),
                "baslangic_tarihi": plan.get("plan_baslangic_tarihi"),
                "aciklama": plan.get("plan_aciklama"),
            }

        cursor.execute(
            """
            SELECT id, sira, vade_tarihi, tutar_cents, durum, odeme_tarihi, aciklama
              FROM odeme_plani_harici
             WHERE harici_finans_id=?
             ORDER BY COALESCE(sira, 0), COALESCE(vade_tarihi,''), id
            """,
            (hid,),
        )
        rows = cursor.fetchall() or []
        installments: list[dict[str, Any]] = []
        for row in rows:
            if isinstance(row, sqlite3.Row):
                data = dict(row)
            else:
                data = {
                    "id": row[0] if len(row) > 0 else None,
                    "sira": row[1] if len(row) > 1 else None,
                    "vade_tarihi": row[2] if len(row) > 2 else None,
                    "tutar_cents": row[3] if len(row) > 3 else 0,
                    "durum": row[4] if len(row) > 4 else "Ödenecek",
                    "odeme_tarihi": row[5] if len(row) > 5 else None,
                    "aciklama": row[6] if len(row) > 6 else "",
                }
            _normalize_installment_status(data)
            installments.append(data)
        return plan, installments
    finally:
        if owns:
            conn.close()


def harici_save_payment_plan(
    conn,
    hid: int,
    plan_data: dict[str, Any],
    taksitler: list[dict[str, Any]],
) -> None:
    """Persist plan metadata and rows for external finance."""

    conn, owns = _ensure_conn(conn)
    try:
        with conn:
            conn.execute(
                """
                UPDATE finans_harici
                   SET plan_taksit_sayisi=?,
                       plan_periyot=?,
                       plan_vade_gunu=?,
                       plan_baslangic_tarihi=?,
                       plan_aciklama=?,
                       updated_at=CURRENT_TIMESTAMP
                 WHERE id=?
                """,
                (
                    int(plan_data.get("taksit_sayisi") or 0),
                    plan_data.get("periyot") or "Ay",
                    int(plan_data.get("vade_gunu") or 1),
                    plan_data.get("baslangic_tarihi"),
                    plan_data.get("aciklama"),
                    hid,
                ),
            )
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id FROM odeme_plani_harici WHERE harici_finans_id=?",
                (hid,),
            )
            existing_ids = {int(row[0]) for row in cursor.fetchall()}
            kept: set[int] = set()
            today = date.today()
            for taksit in taksitler:
                status = _normalize_installment_status(taksit, today=today)
                normalized_status = (status or "Ödenecek").strip()
                values = (
                    taksit.get("sira"),
                    taksit.get("vade_tarihi"),
                    int(taksit.get("tutar_cents") or 0),
                    normalized_status or "Ödenecek",
                    taksit.get("odeme_tarihi"),
                    taksit.get("aciklama"),
                )
                inst_id = taksit.get("id")
                assigned_id: int | None = None
                if inst_id and int(inst_id) in existing_ids:
                    cursor.execute(
                        """
                        UPDATE odeme_plani_harici
                           SET sira=?, vade_tarihi=?, tutar_cents=?, durum=?, odeme_tarihi=?, aciklama=?,
                               updated_at=CURRENT_TIMESTAMP
                         WHERE id=? AND harici_finans_id=?
                        """,
                        (*values, int(inst_id), hid),
                    )
                    assigned_id = int(inst_id)
                    kept.add(assigned_id)
                else:
                    cursor.execute(
                        """
                        INSERT INTO odeme_plani_harici
                          (harici_finans_id, sira, vade_tarihi, tutar_cents, durum, odeme_tarihi, aciklama, created_at, updated_at)
                        VALUES (?,?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                        """,
                        (hid, *values),
                    )
                    inst_id = cursor.lastrowid
                    if inst_id:
                        assigned_id = int(inst_id)
                        kept.add(assigned_id)
                        taksit["id"] = assigned_id
                if (
                    assigned_id
                    and normalized_status.casefold() == "ödendi"
                ):
                    payload = dict(taksit)
                    payload["id"] = assigned_id
                    payload.setdefault("harici_finans_id", hid)
                    try:
                        harici_insert_payment_from_installment(conn, hid, payload)
                    except Exception:
                        # Otomatik ödeme oluşturma hataları ana kaydı engellemesin.
                        pass
            removed = existing_ids - kept
            if removed:
                cursor.executemany(
                    "DELETE FROM odeme_plani_harici WHERE id=?",
                    ((rid,) for rid in removed),
                )
        harici_recalculate_totals(conn, hid)
    finally:
        if owns:
            conn.close()


def insert_payment_from_installment(
    conn, finans_id: int, installment: dict[str, Any] | sqlite3.Row | None
) -> bool:
    """Auto-create a bound-finance payment row for a paid installment."""

    if finans_id in (None, "", 0) or not installment:
        return False

    if isinstance(installment, sqlite3.Row):
        data: dict[str, Any] = dict(installment)
    elif isinstance(installment, dict):
        data = dict(installment)
    else:
        try:
            data = dict(installment)
        except Exception:
            return False

    inst_id = data.get("id") or data.get("taksit_id")
    try:
        inst_id_int = int(inst_id)
    except (TypeError, ValueError):
        return False

    pay_date = _coerce_auto_payment_date(
        data.get("vade_tarihi")
        or data.get("tarih")
        or data.get("odeme_tarihi")
    )
    if not pay_date:
        pay_date = date.today().strftime("%Y-%m-%d")

    try:
        amount_cents = int(data.get("tutar_cents") or 0)
    except (TypeError, ValueError):
        amount_cents = 0

    description = (data.get("aciklama") or "").strip() or AUTO_PAYMENT_NOTE

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT 1 FROM odeme_kayitlari WHERE finans_id=? AND taksit_id=? LIMIT 1",
            (finans_id, inst_id_int),
        )
        if cursor.fetchone():
            return False
        was_in_tx = conn.in_transaction
        cursor.execute(
            """
            INSERT INTO odeme_kayitlari
              (finans_id, tarih, tutar_cents, yontem, aciklama, taksit_id)
            VALUES (?,?,?,?,?,?)
            """,
            (
                finans_id,
                pay_date,
                amount_cents,
                "Taksit",
                description or None,
                inst_id_int,
            ),
        )
        if not was_in_tx:
            conn.commit()
        return True
    finally:
        if owns:
            conn.close()


def delete_payment_by_installment(
    conn, finans_id: int, installment_id: int
) -> bool:
    """Delete a payment row that was auto-created from an installment (bound finance)."""

    if finans_id in (None, "", 0) or installment_id in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        cursor.execute(
            "DELETE FROM odeme_kayitlari WHERE finans_id=? AND taksit_id=?",
            (finans_id, installment_id),
        )
        deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def insert_payment_from_kasadan(
    conn, finans_id: int, tarih: str, tutar_cents: int, aciklama: str | None = None
) -> bool:
    """Kasadan yapılan sözleşme ödemesi için ödeme kaydı oluşturur (bound finance)."""

    if finans_id in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        cursor.execute(
            """
            INSERT INTO odeme_kayitlari
              (finans_id, tarih, tutar_cents, yontem, aciklama)
            VALUES (?,?,?,?,?)
            """,
            (
                finans_id,
                tarih,
                tutar_cents,
                "Kasadan",
                aciklama or "Kasadan ödeme",
            ),
        )
        if not was_in_tx:
            conn.commit()
        return True
    finally:
        if owns:
            conn.close()


def delete_payment_from_kasadan(
    conn, finans_id: int, tarih: str, tutar_cents: int
) -> bool:
    """Kasadan yapılan sözleşme ödemesini siler (bound finance)."""
    if finans_id in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        # Eşleşen ilk kaydı sil (tarih, tutar ve yontem="Kasadan")
        cursor.execute(
            """
            DELETE FROM odeme_kayitlari
            WHERE id = (
                SELECT id FROM odeme_kayitlari
                WHERE finans_id=? AND tarih=? AND tutar_cents=? AND yontem='Kasadan'
                LIMIT 1
            )
            """,
            (finans_id, tarih, tutar_cents),
        )
        deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def delete_payment(payment_id: int) -> bool:
    """Ödeme kaydını ID'ye göre siler."""
    if not payment_id:
        return False
    conn = get_connection()
    try:
        with conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM odeme_kayitlari WHERE id = ?", (payment_id,))
            return cursor.rowcount > 0
    finally:
        conn.close()


def harici_delete_payment(payment_id: int) -> bool:
    """Harici ödeme kaydını ID'ye göre siler."""
    if not payment_id:
        return False
    conn = get_connection()
    try:
        with conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM odemeler_harici WHERE id = ?", (payment_id,))
            return cursor.rowcount > 0
    finally:
        conn.close()


def delete_expense(expense_id: int) -> bool:
    """Masraf kaydını ID'ye göre siler."""
    if not expense_id:
        return False
    conn = get_connection()
    try:
        with conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM masraflar WHERE id = ?", (expense_id,))
            return cursor.rowcount > 0
    finally:
        conn.close()


def harici_delete_expense(expense_id: int) -> bool:
    """Harici masraf kaydını ID'ye göre siler."""
    if not expense_id:
        return False
    conn = get_connection()
    try:
        with conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM masraflar_harici WHERE id = ?", (expense_id,))
            return cursor.rowcount > 0
    finally:
        conn.close()


def delete_expense_by_kasa_avans(
    conn, finans_id: int, tarih: str, tutar_cents: int
) -> bool:
    """Kullanılan Avans silindiğinde ilgili masraf kaydını siler."""
    if finans_id in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        # Önce tarih ve tutar ile eşleşen masrafı bul
        cursor.execute(
            """
            DELETE FROM masraflar
            WHERE id = (
                SELECT id FROM masraflar
                WHERE finans_id=? AND (tarih=? OR tarih IS NULL OR tarih='') AND tutar_cents=? AND odeme_kaynagi='Kasadan'
                LIMIT 1
            )
            """,
            (finans_id, tarih, tutar_cents),
        )
        deleted = cursor.rowcount > 0
        # Eğer tarih ile bulunamadıysa sadece tutar ile dene
        if not deleted:
            cursor.execute(
                """
                DELETE FROM masraflar
                WHERE id = (
                    SELECT id FROM masraflar
                    WHERE finans_id=? AND tutar_cents=? AND odeme_kaynagi='Kasadan'
                    LIMIT 1
                )
                """,
                (finans_id, tutar_cents),
            )
            deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def harici_insert_payment_from_installment(
    conn, hid: int, installment: dict[str, Any] | sqlite3.Row | None
) -> bool:
    """Create a payment row for the provided external installment if needed."""

    if hid in (None, "", 0) or not installment:
        return False

    if isinstance(installment, sqlite3.Row):
        data: dict[str, Any] = dict(installment)
    elif isinstance(installment, dict):
        data = dict(installment)
    else:
        try:
            data = dict(installment)
        except Exception:
            return False

    inst_id = data.get("id")
    try:
        inst_id_int = int(inst_id)
    except (TypeError, ValueError):
        return False

    pay_date = _coerce_auto_payment_date(
        data.get("vade_tarihi")
        or data.get("tarih")
        or data.get("odeme_tarihi")
    )
    if not pay_date:
        pay_date = date.today().strftime("%Y-%m-%d")

    try:
        amount_cents = int(data.get("tutar_cents") or 0)
    except (TypeError, ValueError):
        amount_cents = 0

    description = (data.get("aciklama") or "").strip() or HARICI_AUTO_PAYMENT_NOTE

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT 1 FROM odemeler_harici WHERE harici_finans_id=? AND plan_taksit_id=? LIMIT 1",
            (hid, inst_id_int),
        )
        if cursor.fetchone():
            return False
        was_in_tx = conn.in_transaction
        insert_sql = (
            """
            INSERT INTO odemeler_harici
              (harici_finans_id, tarih, tutar_cents, tahsil_durumu, tahsil_tarihi, yontem, aciklama, plan_taksit_id, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?, ?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
            """
        )
        params = (
            hid,
            pay_date,
            amount_cents,
            "Ödendi",
            pay_date,
            "Taksit",
            description or None,
            inst_id_int,
        )
        try:
            cursor.execute(insert_sql, params)
        except sqlite3.OperationalError:
            cursor.execute(
                """
                INSERT INTO odemeler_harici
                  (harici_finans_id, tahsil_tarihi, tutar_cents, yontem, aciklama, created_at, updated_at)
                VALUES (?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                """,
                (hid, pay_date, amount_cents, "Taksit", description or None),
            )
        if not was_in_tx:
            conn.commit()
        return True
    finally:
        if owns:
            conn.close()


def harici_delete_payment_by_installment(
    conn, hid: int, installment_id: int
) -> bool:
    """Delete a payment row that was auto-created from an installment."""

    if hid in (None, "", 0) or installment_id in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        cursor.execute(
            "DELETE FROM odemeler_harici WHERE harici_finans_id=? AND plan_taksit_id=?",
            (hid, installment_id),
        )
        deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def harici_insert_payment_from_kasadan(
    conn, hid: int, tarih: str, tutar_cents: int, aciklama: str | None = None
) -> bool:
    """Kasadan yapılan sözleşme ödemesi için ödeme kaydı oluşturur."""

    if hid in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        insert_sql = """
            INSERT INTO odemeler_harici
              (harici_finans_id, tarih, tutar_cents, tahsil_durumu, tahsil_tarihi, yontem, aciklama, created_at, updated_at)
            VALUES (?,?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
        """
        params = (
            hid,
            tarih,
            tutar_cents,
            "Ödendi",
            tarih,
            "Kasadan",
            aciklama or "Kasadan ödeme",
        )
        try:
            cursor.execute(insert_sql, params)
        except sqlite3.OperationalError:
            cursor.execute(
                """
                INSERT INTO odemeler_harici
                  (harici_finans_id, tahsil_tarihi, tutar_cents, yontem, aciklama, created_at, updated_at)
                VALUES (?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                """,
                (hid, tarih, tutar_cents, "Kasadan", aciklama or "Kasadan ödeme"),
            )
        if not was_in_tx:
            conn.commit()
        return True
    finally:
        if owns:
            conn.close()


def harici_delete_payment_from_kasadan(
    conn, hid: int, tarih: str, tutar_cents: int
) -> bool:
    """Kasadan yapılan sözleşme ödemesini siler (harici finans)."""
    if hid in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        # Eşleşen ilk kaydı sil (tarih, tutar ve yontem="Kasadan")
        cursor.execute(
            """
            DELETE FROM odemeler_harici
            WHERE id = (
                SELECT id FROM odemeler_harici
                WHERE harici_finans_id=? AND (tarih=? OR tahsil_tarihi=?) AND tutar_cents=? AND yontem='Kasadan'
                LIMIT 1
            )
            """,
            (hid, tarih, tarih, tutar_cents),
        )
        deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def harici_delete_expense_by_kasa_avans(
    conn, hid: int, tarih: str, tutar_cents: int
) -> bool:
    """Kullanılan Avans silindiğinde ilgili masraf kaydını siler (harici finans)."""
    if hid in (None, "", 0):
        return False

    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        was_in_tx = conn.in_transaction
        # Önce tarih ve tutar ile eşleşen masrafı bul
        cursor.execute(
            """
            DELETE FROM masraflar_harici
            WHERE id = (
                SELECT id FROM masraflar_harici
                WHERE harici_finans_id=? AND (tarih=? OR tarih IS NULL OR tarih='') AND tutar_cents=? AND odeme_kaynagi='Kasadan'
                LIMIT 1
            )
            """,
            (hid, tarih, tutar_cents),
        )
        deleted = cursor.rowcount > 0
        # Eğer tarih ile bulunamadıysa sadece tutar ile dene
        if not deleted:
            cursor.execute(
                """
                DELETE FROM masraflar_harici
                WHERE id = (
                    SELECT id FROM masraflar_harici
                    WHERE harici_finans_id=? AND tutar_cents=? AND odeme_kaynagi='Kasadan'
                    LIMIT 1
                )
                """,
                (hid, tutar_cents),
            )
            deleted = cursor.rowcount > 0
        if not was_in_tx:
            conn.commit()
        return deleted
    finally:
        if owns:
            conn.close()


def harici_get_payments(conn, hid: int) -> list[dict[str, Any]]:
    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, tarih, COALESCE(tutar_cents,0) AS tutar_cents,
                   yontem, aciklama, tahsil_durumu, tahsil_tarihi, plan_taksit_id
              FROM odemeler_harici
             WHERE harici_finans_id=?
             ORDER BY COALESCE(tahsil_tarihi,''), id
            """,
            (hid,),
        )
        rows = cursor.fetchall() or []
        payments: list[dict[str, Any]] = []
        for row in rows:
            if isinstance(row, sqlite3.Row):
                payments.append(dict(row))
            else:
                payments.append(
                    {
                        "id": row[0] if len(row) > 0 else None,
                        "tarih": row[1] if len(row) > 1 else None,
                        "tutar_cents": int(row[2] if len(row) > 2 else 0),
                        "yontem": row[3] if len(row) > 3 else "",
                        "aciklama": row[4] if len(row) > 4 else "",
                        "tahsil_durumu": row[5] if len(row) > 5 else "Ödendi",
                        "tahsil_tarihi": row[6] if len(row) > 6 else None,
                        "plan_taksit_id": row[7] if len(row) > 7 else None,
                    }
                )
        return payments
    finally:
        if owns:
            conn.close()


def harici_save_payments(conn, hid: int, payments: list[dict[str, Any]]) -> None:
    conn, owns = _ensure_conn(conn)
    try:
        with conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id FROM odemeler_harici WHERE harici_finans_id=?",
                (hid,),
            )
            existing = {int(row[0]) for row in cursor.fetchall()}
            kept: set[int] = set()
            total_cents = 0
            for record in payments:
                amount = int(record.get("tutar_cents") or 0)
                total_cents += max(0, amount)
                tarih_value = record.get("tarih") or record.get("tahsil_tarihi")
                tahsil_value = record.get("tahsil_tarihi") or tarih_value
                status_value = record.get("tahsil_durumu") or "Ödendi"
                plan_taksit = record.get("plan_taksit_id")
                if plan_taksit in ("", 0):
                    plan_taksit = None
                values = (
                    tarih_value,
                    tahsil_value,
                    amount,
                    record.get("yontem"),
                    record.get("aciklama"),
                    status_value,
                    plan_taksit,
                )
                rec_id = record.get("id")
                if rec_id and int(rec_id) in existing:
                    cursor.execute(
                        """
                        UPDATE odemeler_harici
                           SET tarih=?, tahsil_tarihi=?, tutar_cents=?, yontem=?, aciklama=?,
                               tahsil_durumu=?, plan_taksit_id=?,
                               updated_at=CURRENT_TIMESTAMP
                         WHERE id=? AND harici_finans_id=?
                        """,
                        (*values, int(rec_id), hid),
                    )
                    kept.add(int(rec_id))
                else:
                    cursor.execute(
                        """
                        INSERT INTO odemeler_harici
                          (harici_finans_id, tarih, tahsil_tarihi, tutar_cents, yontem, aciklama, tahsil_durumu, plan_taksit_id, created_at, updated_at)
                        VALUES (?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                        """,
                        (hid, *values),
                    )
                    rec_id = cursor.lastrowid
                    if rec_id:
                        kept.add(int(rec_id))
                        record["id"] = int(rec_id)
            removed = existing - kept
            if removed:
                cursor.executemany(
                    "DELETE FROM odemeler_harici WHERE id=?",
                    ((rid,) for rid in removed),
                )
            cursor.execute(
                """
                UPDATE finans_harici
                   SET tahsil_edilen_cents=?, updated_at=CURRENT_TIMESTAMP
                 WHERE id=?
                """,
                (total_cents, hid),
            )
        harici_recalculate_totals(conn, hid)
    finally:
        if owns:
            conn.close()


def harici_get_expenses(conn, hid: int) -> list[dict[str, Any]]:
    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()
        cursor.execute(
            """
            SELECT id, kalem, COALESCE(tutar_cents,0) AS tutar_cents,
                   tarih, tahsil_durumu, tahsil_tarihi, aciklama,
                   COALESCE(odeme_kaynagi, 'Büro') AS odeme_kaynagi
              FROM masraflar_harici
             WHERE harici_finans_id=?
             ORDER BY id ASC
            """,
            (hid,),
        )
        rows = cursor.fetchall() or []
        expenses: list[dict[str, Any]] = []
        for row in rows:
            if isinstance(row, sqlite3.Row):
                expenses.append(dict(row))
            else:
                expenses.append(
                    {
                        "id": row[0] if len(row) > 0 else None,
                        "kalem": row[1] if len(row) > 1 else "",
                        "tutar_cents": int(row[2] if len(row) > 2 else 0),
                        "tarih": row[3] if len(row) > 3 else None,
                        "tahsil_durumu": row[4] if len(row) > 4 else "Bekliyor",
                        "tahsil_tarihi": row[5] if len(row) > 5 else None,
                        "aciklama": row[6] if len(row) > 6 else "",
                        "odeme_kaynagi": row[7] if len(row) > 7 else "Büro",
                    }
                )
        return expenses
    finally:
        if owns:
            conn.close()


def harici_save_expenses(conn, hid: int, rows: list[dict[str, Any]]) -> None:
    conn, owns = _ensure_conn(conn)
    try:
        cursor = conn.cursor()

        # Eski masrafları al (silme kontrolü için)
        cursor.execute(
            "SELECT id FROM masraflar_harici WHERE harici_finans_id=?",
            (hid,),
        )
        existing = {int(row[0]) for row in cursor.fetchall()}

        with conn:
            # Önce tüm eski "Kullanılan Avans" kayıtlarını sil
            cursor.execute(
                "DELETE FROM harici_muvekkil_kasasi WHERE harici_finans_id = ? AND islem_turu = 'Kullanılan Avans'",
                (hid,),
            )

            kept: set[int] = set()
            total_c = 0
            collected_c = 0

            for record in rows:
                amount = int(record.get("tutar_cents") or 0)
                odeme_kaynagi = record.get("odeme_kaynagi") or "Büro"
                tarih = record.get("tarih")
                kalem = record.get("kalem") or ""

                # Sadece Büro masrafları büronun alacağı olarak sayılır
                if odeme_kaynagi == "Büro":
                    total_c += max(0, amount)
                status = record.get("tahsil_durumu") or "Bekliyor"
                tahsil_amount = amount if status == "Tahsil Edildi" else 0
                if odeme_kaynagi == "Büro":
                    collected_c += max(0, tahsil_amount)

                values = (
                    kalem,
                    amount,
                    tarih,
                    status,
                    record.get("tahsil_tarihi"),
                    tahsil_amount,
                    record.get("aciklama"),
                    odeme_kaynagi,
                )

                rec_id = record.get("id")
                if rec_id and int(rec_id) in existing:
                    cursor.execute(
                        """
                        UPDATE masraflar_harici
                           SET kalem=?, tutar_cents=?, tarih=?, tahsil_durumu=?, tahsil_tarihi=?,
                               tahsil_cents=?, aciklama=?, odeme_kaynagi=?, updated_at=CURRENT_TIMESTAMP
                         WHERE id=? AND harici_finans_id=?
                        """,
                        (*values, int(rec_id), hid),
                    )
                    kept.add(int(rec_id))
                else:
                    cursor.execute(
                        """
                        INSERT INTO masraflar_harici
                          (harici_finans_id, kalem, tutar_cents, tarih, tahsil_durumu, tahsil_tarihi, tahsil_cents, aciklama, odeme_kaynagi, created_at, updated_at)
                        VALUES (?,?,?,?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                        """,
                        (hid, *values),
                    )
                    rec_id = cursor.lastrowid
                    if rec_id:
                        kept.add(int(rec_id))
                        record["id"] = int(rec_id)

                # Kasadan masraf için müvekkil kasasına "Kullanılan Avans" ekle
                if odeme_kaynagi == "Kasadan" and amount > 0:
                    kasa_tarih = tarih or date.today().isoformat()
                    cursor.execute(
                        """
                        INSERT INTO harici_muvekkil_kasasi
                          (harici_finans_id, tarih, tutar_kurus, islem_turu, aciklama)
                        VALUES (?, ?, ?, 'Kullanılan Avans', ?)
                        """,
                        (hid, kasa_tarih, amount, f"Masraf: {kalem}"),
                    )

            removed = existing - kept
            if removed:
                cursor.executemany(
                    "DELETE FROM masraflar_harici WHERE id=?",
                    ((rid,) for rid in removed),
                )
            cursor.execute(
                """
                UPDATE finans_harici
                   SET masraf_toplam_cents=?, masraf_tahsil_cents=?, updated_at=CURRENT_TIMESTAMP
                 WHERE id=?
                """,
                (total_c, collected_c, hid),
            )
        harici_recalculate_totals(conn, hid)
    finally:
        if owns:
            conn.close()


def harici_recalculate_totals(conn, hid: int) -> None:
    """Recompute stored summary totals for an external finance record."""

    conn, owns = _ensure_conn(conn)
    try:
        with conn:
            cursor = conn.cursor()
            today = date.today()
            cursor.execute(
                "SELECT id, vade_tarihi, durum FROM odeme_plani_harici WHERE harici_finans_id=?",
                (hid,),
            )
            installments = cursor.fetchall() or []
            has_overdue = 0
            for row in installments:
                if isinstance(row, sqlite3.Row):
                    inst_id = row["id"] if "id" in row.keys() else row[0]
                    due_date = row["vade_tarihi"] if "vade_tarihi" in row.keys() else row[1]
                    prev_status = row["durum"] if "durum" in row.keys() else row[2]
                else:
                    inst_id = row[0] if len(row) > 0 else None
                    due_date = row[1] if len(row) > 1 else None
                    prev_status = row[2] if len(row) > 2 else "Ödenecek"
                prev_status = prev_status or "Ödenecek"
                inst_dict = {"id": inst_id, "vade_tarihi": due_date, "durum": prev_status}
                updated_status = _normalize_installment_status(inst_dict, today=today)
                if updated_status == INSTALLMENT_OVERDUE_STATUS:
                    has_overdue = 1
                if inst_id and updated_status != prev_status:
                    cursor.execute(
                        "UPDATE odeme_plani_harici SET durum=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",
                        (updated_status, int(inst_id)),
                    )

            cursor.execute(
                """
                SELECT
                    COALESCE(sabit_ucret_cents, 0) AS sabit_ucret_cents,
                    COALESCE(yuzde_orani, 0) AS yuzde_orani,
                    COALESCE(tahsil_hedef_cents, 0) AS tahsil_hedef_cents,
                    COALESCE(yuzde_is_sonu, 0) AS yuzde_is_sonu
                  FROM finans_harici
                 WHERE id=?
                """,
                (hid,),
            )
            row = cursor.fetchone()
            if not row:
                return
            if isinstance(row, sqlite3.Row):
                contract_row = dict(row)
            else:
                contract_row = {
                    "sabit_ucret_cents": row[0] if len(row) > 0 else 0,
                    "yuzde_orani": row[1] if len(row) > 1 else 0,
                    "tahsil_hedef_cents": row[2] if len(row) > 2 else 0,
                    "yuzde_is_sonu": row[3] if len(row) > 3 else 0,
                }

            total = calculate_harici_total(contract_row)
            fixed = int(contract_row.get("sabit_ucret_cents") or 0)

            cursor.execute(
                "SELECT COALESCE(SUM(tutar_cents), 0) FROM odemeler_harici WHERE harici_finans_id=?",
                (hid,),
            )
            collected = int(cursor.fetchone()[0] or 0)
            # Sadece Büro masrafları büronun alacağı olarak sayılır
            cursor.execute(
                """
                SELECT
                    COALESCE(SUM(tutar_cents), 0) AS toplam,
                    COALESCE(
                        SUM(CASE WHEN tahsil_durumu = 'Tahsil Edildi' THEN tutar_cents ELSE 0 END),
                        0
                    ) AS tahsil
                  FROM masraflar_harici
                 WHERE harici_finans_id=?
                   AND (odeme_kaynagi IS NULL OR odeme_kaynagi = 'Büro')
                """,
                (hid,),
            )
            expense_row = cursor.fetchone() or (0, 0)
            if isinstance(expense_row, sqlite3.Row):
                expense = int(expense_row["toplam"] if "toplam" in expense_row.keys() else expense_row[0])
                expense_collected = int(
                    expense_row["tahsil"] if "tahsil" in expense_row.keys() else expense_row[1]
                )
            else:
                expense = int(expense_row[0] if len(expense_row) > 0 else 0)
                expense_collected = int(expense_row[1] if len(expense_row) > 1 else 0)

            contract_row.update(
                {
                    "tahsil_edilen_cents": collected,
                    "masraf_toplam_cents": expense,
                    "masraf_tahsil_cents": expense_collected,
                }
            )
            balance = calculate_harici_balance(contract_row)
            cursor.execute(
                """
                UPDATE finans_harici
                   SET sabit_ucret_cents=?,
                       tahsil_edilen_cents=?,
                       masraf_toplam_cents=?,
                       masraf_tahsil_cents=?,
                       toplam_ucret_cents=?,
                       kalan_bakiye_cents=?,
                       has_overdue_installment=?,
                       updated_at=CURRENT_TIMESTAMP
                 WHERE id=?
                """,
                (
                    fixed,
                    collected,
                    expense,
                    expense_collected,
                    total,
                    balance,
                    has_overdue,
                    hid,
                ),
            )
    finally:
        if owns:
            conn.close()
def harici_masraflar_load(conn, hid: int):
    if conn is None:
        return []
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT id, tarih, kalem, COALESCE(tutar_cents,0) AS tutar_c,
               COALESCE(tahsil_cents,0) AS tahsil_c, aciklama
        FROM masraflar_harici
        WHERE harici_finans_id=?
        ORDER BY id ASC
        """,
        (hid,),
    )
    return cursor.fetchall()


def harici_masraflar_replace(conn, hid: int, items: list[dict]):
    if conn is None:
        return
    with conn:
        conn.execute(
            "DELETE FROM masraflar_harici WHERE harici_finans_id=?",
            (hid,),
        )
        for item in items:
            conn.execute(
                """
                INSERT INTO masraflar_harici
                  (harici_finans_id, tarih, kalem, tutar_cents, tahsil_cents, aciklama, created_at, updated_at)
                VALUES (?,?,?,?,?,?,CURRENT_TIMESTAMP,CURRENT_TIMESTAMP)
                """,
                (
                    hid,
                    item.get("tarih_iso"),
                    item.get("kalem") or "",
                    item.get("tutar_c") or 0,
                    item.get("tahsil_c") or 0,
                    item.get("aciklama") or "",
                ),
            )


def harici_masraflar_sumlar(conn, hid: int) -> tuple[int, int]:
    """Calculate total expense and collection cents for a record."""

    if conn is None:
        return (0, 0)
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT COALESCE(SUM(tutar_cents),0) AS t, COALESCE(SUM(tahsil_cents),0) AS th
          FROM masraflar_harici
         WHERE harici_finans_id=?
        """,
        (hid,),
    )
    row = cursor.fetchone()
    if not row:
        return (0, 0)
    try:
        toplam = row["t"]
    except (KeyError, IndexError, TypeError):
        toplam = row[0] if row else 0
    try:
        tahsil = row["th"]
    except (KeyError, IndexError, TypeError):
        tahsil = row[1] if row else 0
    return int(toplam or 0), int(tahsil or 0)


def harici_update_masraf_ozet(
    conn, hid: int, toplam_c: int, tahsil_c: int
) -> None:
    """Persist aggregate expense information on the parent record."""

    if conn is None:
        return
    with conn:
        conn.execute(
            """
            UPDATE finans_harici
               SET masraf_toplam_cents=?, masraf_tahsil_cents=?, updated_at=CURRENT_TIMESTAMP
             WHERE id=?
            """,
            (toplam_c or 0, tahsil_c or 0, hid),
        )


def harici_odemeler_list(conn, hid: int):
    """Return payment rows for an external finance record ordered by date."""

    if conn is None:
        return []
    cursor = conn.cursor()
    cursor.execute(
        """
        SELECT tahsil_tarihi, COALESCE(tutar_cents,0) AS tutar_c
          FROM odemeler_harici
         WHERE harici_finans_id=?
         ORDER BY COALESCE(tahsil_tarihi,''), id
        """,
        (hid,),
    )
    rows = cursor.fetchall() or []
    results: list[dict[str, Any]] = []
    for row in rows:
        if isinstance(row, dict):
            tarih = row.get("tahsil_tarihi") or ""
            tutar = row.get("tutar_c") or 0
        else:
            try:
                tarih = row[0]
            except (IndexError, TypeError):
                tarih = ""
            try:
                tutar = row[1]
            except (IndexError, TypeError):
                tutar = 0
        results.append({"tarih": tarih or "", "tutar_c": int(tutar or 0)})
    return results


def harici_create(conn: sqlite3.Connection | None = None) -> int:
    owns_conn = False
    if conn is None:
        conn = get_connection()
        owns_conn = True
    try:
        cur = conn.cursor()
        cur.execute(
            """
            INSERT INTO finans_harici (created_at, updated_at)
            VALUES (CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
            """
        )
        if owns_conn:
            conn.commit()
        return int(cur.lastrowid)
    finally:
        if owns_conn:
            conn.close()


def harici_update_quick_info(
    conn: sqlite3.Connection | None, harici_id: int, bn: str, muv: str
) -> None:
    owns_conn = False
    if conn is None:
        conn = get_connection()
        owns_conn = True
    try:
        with conn:
            conn.execute(
                """
                UPDATE finans_harici
                   SET harici_bn = ?,
                       harici_muvekkil = ?,
                       updated_at = CURRENT_TIMESTAMP
                 WHERE id = ?
                """,
                (bn, muv, harici_id),
            )
    finally:
        if owns_conn and conn is not None:
            conn.close()


def harici_get_master_list(
    conn: sqlite3.Connection | None = None, *, search_text: str = ""
) -> list[dict]:
    owns_conn = False
    if conn is None:
        conn = get_connection()
        owns_conn = True
    try:
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute(
            """
            WITH next_due AS (
                SELECT harici_finans_id AS finans_id, MIN(vade_tarihi) AS next_due_date
                  FROM odeme_plani_harici
                 WHERE COALESCE(TRIM(durum), 'Ödenecek') != 'Ödendi'
                 GROUP BY harici_finans_id
            ),
            overdue AS (
                SELECT harici_finans_id
                  FROM odeme_plani_harici
                 WHERE COALESCE(TRIM(durum), 'Ödenecek') != 'Ödendi'
                   AND COALESCE(vade_tarihi, '') != ''
                   AND DATE(vade_tarihi) < DATE('now')
                 GROUP BY harici_finans_id
            )
            SELECT
                fh.id,
                fh.harici_bn,
                fh.harici_esas_no,
                fh.harici_muvekkil,
                COALESCE(fh.sabit_ucret_cents, 0) AS sabit_ucret_cents,
                COALESCE(fh.yuzde_orani, 0) AS yuzde_orani,
                COALESCE(fh.tahsil_hedef_cents, 0) AS tahsil_hedef_cents,
                COALESCE(fh.yuzde_is_sonu, 0) AS yuzde_is_sonu,
                COALESCE(fh.tahsil_edilen_cents, 0) AS tahsil_edilen_cents,
                COALESCE(fh.masraf_toplam_cents, 0) AS masraf_toplam_cents,
                COALESCE(fh.masraf_tahsil_cents, 0) AS masraf_tahsil_cents,
                COALESCE(fh.toplam_ucret_cents, 0) AS toplam_ucret_cents,
                COALESCE(fh.kalan_bakiye_cents, 0) AS kalan_bakiye_cents,
                fh.created_at,
                fh.updated_at,
                next_due.next_due_date,
                CASE
                    WHEN overdue.harici_finans_id IS NOT NULL THEN 1
                    ELSE COALESCE(fh.has_overdue_installment, 0)
                END AS has_overdue_installment
            FROM finans_harici fh
            LEFT JOIN next_due ON next_due.finans_id = fh.id
            LEFT JOIN overdue ON overdue.harici_finans_id = fh.id
            ORDER BY fh.id DESC
            """
        )
        rows = cur.fetchall() or []
    finally:
        if owns_conn:
            conn.close()

    processed: list[dict] = []
    for row in rows:
        record = dict(row)
        record["id"] = int(record.get("id") or 0)
        record["harici_bn"] = record.get("harici_bn") or ""
        record["harici_esas_no"] = record.get("harici_esas_no") or ""
        record["harici_muvekkil"] = record.get("harici_muvekkil") or ""
        record["assigned_user_ids"] = []
        total = int(record.get("toplam_ucret_cents") or 0)
        if total <= 0:
            total = calculate_harici_total(record)
        record["toplam_ucret_cents"] = total
        record["kalan_bakiye_cents"] = calculate_harici_balance(record)
        record["has_overdue_installment"] = bool(record.get("has_overdue_installment"))
        record["due_category"] = _categorize_due_date(record.get("next_due_date"))
        processed.append(record)

    search = (search_text or "").strip().lower()
    if not search:
        return processed

    filtered: list[dict] = []
    for record in processed:
        bn = record.get("harici_bn")
        esas = record.get("harici_esas_no")
        muv = record.get("harici_muvekkil")
        if any(
            value and search in str(value).lower()
            for value in (bn, esas, muv)
        ):
            filtered.append(record)
    return filtered


def _categorize_due_date(next_due: Optional[str]) -> Optional[str]:
    if not next_due:
        return None
    parsed_date: Optional[date] = None
    for fmt in ("%Y-%m-%d", "%d.%m.%Y"):
        try:
            parsed_date = datetime.strptime(next_due, fmt).date()
            break
        except ValueError:
            continue
    if parsed_date is None:
        return None
    today = date.today()
    days_left = (parsed_date - today).days
    if days_left < 0:
        return "overdue"
    if days_left == 0:
        return "due_today"
    if 1 <= days_left <= 3:
        return "due_1_3"
    if 4 <= days_left <= 7:
        return "due_4_7"
    return "due_future"


def update_finans_contract(
    dosya_id: int | None,
    *,
    finans_id: int | None = None,
    sozlesme_ucreti: float | None,
    sozlesme_yuzdesi: float | None,
    tahsil_hedef_cents: int,
    notlar: str | None,
    yuzde_is_sonu: bool,
    karsi_vekalet_ucreti_cents: int = 0,
) -> bool:
    """Finans sözleşme bilgilerini günceller.

    ``True`` dönerse en az bir satır güncellenmiştir; ``False`` dönerse
    eşleşen kayıt bulunamamıştır.
    """

    conn = get_connection()
    try:
        cur = conn.cursor()
        finans_key: int | None = finans_id
        where_clause: str
        where_value: int
        if finans_key is None:
            if dosya_id is None:
                raise ValueError("Finans kaydı güncellenemedi: kimlik belirtilmedi.")
            finans_key = ensure_finans_record(dosya_id, cur)
            where_clause = "dosya_id = ?"
            where_value = dosya_id
        else:
            where_clause = "id = ?"
            where_value = finans_key
        fixed_value = None if sozlesme_ucreti is None else float(sozlesme_ucreti)
        fixed_cents = (
            tl_to_cents(fixed_value)
            if fixed_value is not None
            else None
        )
        percent_value = None if sozlesme_yuzdesi is None else float(sozlesme_yuzdesi)
        target_value = int(tahsil_hedef_cents or 0)
        sql = f"""
            UPDATE finans
            SET sozlesme_ucreti = ?,
                sozlesme_ucreti_cents = ?,
                sozlesme_yuzdesi = ?,
                tahsil_hedef_cents = ?,
                notlar = ?,
                yuzde_is_sonu = ?,
                karsi_vekalet_ucreti_cents = ?,
                son_guncelleme = ?
            WHERE {where_clause}
        """
        params = (
            fixed_value,
            fixed_cents,
            percent_value,
            target_value,
            notlar,
            1 if yuzde_is_sonu else 0,
            int(karsi_vekalet_ucreti_cents or 0),
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            where_value,
        )
        logger.info(
            "Finans kaydı güncelleniyor",
            extra={
                "dosya_id": dosya_id,
                "finans_id": finans_key,
                "sql": " ".join(sql.split()),
                "params": params,
            },
        )
        cur.execute(sql, params)
        updated = cur.rowcount
        if updated and finans_key is not None:
            recalculate_finans_totals(finans_key, cur)
        conn.commit()
        if updated:
            logger.info(
                "Finans kaydı güncellendi",
                extra={
                    "dosya_id": dosya_id,
                    "finans_id": finans_key,
                    "updated_rows": updated,
                },
            )
        else:
            logger.warning(
                "Finans kaydı güncellenemedi (eşleşen satır yok)",
                extra={
                    "dosya_id": dosya_id,
                    "finans_id": finans_key,
                    "sql": " ".join(sql.split()),
                    "params": params,
                },
            )
        return updated > 0
    except sqlite3.Error as exc:
        conn.rollback()
        logger.exception(
            "Finans bilgileri güncellenemedi",
            extra={
                "dosya_id": dosya_id,
                "finans_id": finans_key,
                "sql": " ".join(sql.split()),
                "params": params,
            },
        )
        raise RuntimeError("Finans bilgileri güncellenemedi.") from exc
    finally:
        conn.close()


def update_finans_terms(
    dosya_id: int | None,
    *,
    finans_id: int | None = None,
    sozlesme_ucreti: float | None,
    sozlesme_yuzdesi: float | None,
) -> bool:
    """Sözleşme tutarı ve yüzde alanlarını tek başına günceller.

    ``True`` dönerse en az bir satır güncellenmiştir; ``False`` dönerse
    eşleşen kayıt bulunamamıştır.
    """

    conn = get_connection()
    sql = ""
    params: tuple[Any, ...] | None = None
    try:
        cur = conn.cursor()
        where_clause: str
        where_value: int
        if finans_id is not None:
            where_clause = "id = ?"
            where_value = finans_id
        else:
            if dosya_id is None:
                raise ValueError("Finans sözleşmesi güncellenemedi: kimlik yok.")
            ensure_finans_record(dosya_id, cur)
            where_clause = "dosya_id = ?"
            where_value = dosya_id
        fixed_value = None if sozlesme_ucreti is None else float(sozlesme_ucreti)
        fixed_cents = (
            tl_to_cents(fixed_value) if fixed_value is not None else None
        )
        percent_value = None if sozlesme_yuzdesi is None else float(sozlesme_yuzdesi)
        sql = f"""
            UPDATE finans
            SET sozlesme_ucreti = ?,
                sozlesme_ucreti_cents = ?,
                sozlesme_yuzdesi = ?,
                son_guncelleme = ?
            WHERE {where_clause}
        """
        params = (
            fixed_value,
            fixed_cents,
            percent_value,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            where_value,
        )
        logger.info(
            "Finans sözleşme şartları güncelleniyor",
            extra={
                "dosya_id": dosya_id,
                "finans_id": finans_id,
                "sql": " ".join(sql.split()),
                "params": params,
            },
        )
        cur.execute(sql, params)
        updated = cur.rowcount
        conn.commit()
        if updated:
            logger.info(
                "Finans sözleşme şartları güncellendi",
                extra={
                    "dosya_id": dosya_id,
                    "finans_id": finans_id,
                    "updated_rows": updated,
                },
            )
        else:
            logger.warning(
                "Finans sözleşme şartları güncellenemedi (eşleşen satır yok)",
                extra={
                    "dosya_id": dosya_id,
                    "finans_id": finans_id,
                    "sql": " ".join(sql.split()),
                    "params": params,
                },
            )
        return updated > 0
    except sqlite3.Error as exc:
        conn.rollback()
        logger.exception(
            "Sözleşme bilgileri güncellenemedi",
            extra={
                "dosya_id": dosya_id,
                "finans_id": finans_id,
                "sql": " ".join(sql.split()),
                "params": params,
            },
        )
        raise RuntimeError("Sözleşme bilgileri güncellenemedi.") from exc
    finally:
        conn.close()


def _touch_finans(cur: sqlite3.Cursor, finans_id: int) -> None:
    cur.execute(
        "UPDATE finans SET son_guncelleme = ? WHERE id = ?",
        (datetime.now().strftime("%Y-%m-%d %H:%M:%S"), finans_id),
    )


def recalculate_finans_totals(
    finans_id: int, cur: sqlite3.Cursor | None = None
) -> None:
    """Ödeme ve masraf toplamlarını yeniden hesaplar."""

    owns_connection = cur is None
    conn: sqlite3.Connection | None = None
    if owns_connection:
        conn = get_connection()
        cur = conn.cursor()
    assert cur is not None
    cur.execute(
        "SELECT COALESCE(SUM(tutar_cents), 0) FROM odeme_kayitlari WHERE finans_id = ?",
        (finans_id,),
    )
    tahsil = int(cur.fetchone()[0] or 0)

    # Sadece "Büro" masrafları büronun alacağı olarak sayılır
    # "Kasadan" masraflar müvekkilin avansından ödenmiş, bizim alacağımız değil
    cur.execute(
        """
        SELECT COALESCE(SUM(tutar_cents), 0)
        FROM masraflar
        WHERE finans_id = ? AND (odeme_kaynagi IS NULL OR odeme_kaynagi = 'Büro')
        """,
        (finans_id,),
    )
    masraf_toplam = int(cur.fetchone()[0] or 0)

    cur.execute(
        """
        SELECT COALESCE(SUM(tutar_cents), 0)
        FROM masraflar
        WHERE finans_id = ?
          AND (odeme_kaynagi IS NULL OR odeme_kaynagi = 'Büro')
          AND tahsil_durumu = 'Tahsil Edildi'
        """,
        (finans_id,),
    )
    masraf_tahsil = int(cur.fetchone()[0] or 0)

    cur.execute(
        """
        UPDATE finans
        SET tahsil_edilen_cents = ?,
            masraf_toplam_cents = ?,
            masraf_tahsil_cents = ?,
            son_guncelleme = ?
        WHERE id = ?
        """,
        (
            tahsil,
            masraf_toplam,
            masraf_tahsil,
            datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            finans_id,
        ),
    )
    if owns_connection and conn is not None:
        conn.commit()
        conn.close()


def delete_finans_record(finans_id: int) -> None:
    """Delete a finance record by its identifier."""

    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM finans WHERE id = ?", (finans_id,))
        conn.commit()
    finally:
        conn.close()


def mark_next_installment_paid(finans_id: int) -> Optional[Dict[str, Any]]:
    """İlk ödenecek taksiti ödenmiş olarak işaretler."""

    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT id, tutar_cents
            FROM taksitler
            WHERE finans_id = ? AND durum = 'Ödenecek'
            ORDER BY vade_tarihi, id
            LIMIT 1
            """,
            (finans_id,),
        )
        row = cur.fetchone()
        if row is None:
            return None
        taksit_id, tutar_cents = int(row[0]), int(row[1] or 0)
        today = datetime.now().strftime("%Y-%m-%d")
        cur.execute(
            """
            UPDATE taksitler
            SET durum = 'Ödendi', odeme_tarihi = ?
            WHERE id = ?
            """,
            (today, taksit_id),
        )
        if tutar_cents > 0:
            cur.execute(
                """
                INSERT INTO odeme_kayitlari (finans_id, tarih, tutar_cents, yontem, aciklama)
                VALUES (?, ?, ?, ?, ?)
                """,
                (
                    finans_id,
                    today,
                    tutar_cents,
                    "Taksit",
                    AUTO_PAYMENT_NOTE,
                ),
            )
        recalculate_finans_totals(finans_id, cur)
        conn.commit()
        return {
            "taksit_id": taksit_id,
            "tutar_cents": tutar_cents,
            "odeme_tarihi": today,
        }
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Taksit ödenmiş olarak işaretlenemedi.") from exc
    finally:
        conn.close()


def add_partial_payment(
    finans_id: int,
    tarih: str,
    tutar_cents: int,
    yontem: Optional[str] = None,
    aciklama: Optional[str] = None,
) -> None:
    """Finans kaydına yeni bir ödeme hareketi ekler."""

    if tutar_cents <= 0:
        raise ValueError("Tutar sıfırdan büyük olmalıdır.")
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute(
            """
            INSERT INTO odeme_kayitlari (finans_id, tarih, tutar_cents, yontem, aciklama)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                finans_id,
                tarih,
                tutar_cents,
                yontem,
                aciklama,
            ),
        )
        recalculate_finans_totals(finans_id, cur)
        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Kısmi ödeme kaydedilemedi.") from exc
    finally:
        conn.close()



def _add_months(base_date: date, months: int, vade_gunu: int) -> date:
    year = base_date.year + ((base_date.month - 1 + months) // 12)
    month = ((base_date.month - 1 + months) % 12) + 1
    normalized_day = max(1, min(28, int(vade_gunu or 1)))
    day = min(normalized_day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def generate_installments(
    finans_id: int,
    taksit_sayisi: int,
    periyot: str,
    vade_gunu: int,
    baslangic_tarihi: date | None,
) -> List[Dict[str, Any]]:
    """Finans kaydı için taksit listesini hesaplar."""

    if taksit_sayisi <= 0:
        return []
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM finans WHERE id = ?", (finans_id,))
    finans_row = cur.fetchone()
    conn.close()
    if not finans_row:
        return []
    percent_flag = int(_row_value(finans_row, "yuzde_is_sonu") or 0)

    fixed_cents = int(_row_value(finans_row, "sozlesme_ucreti_cents") or 0)
    if fixed_cents == 0:
        fixed_raw = _row_value(finans_row, "sozlesme_ucreti")
        if fixed_raw not in (None, ""):
            try:
                fixed_cents = tl_to_cents(fixed_raw)
            except (TypeError, ValueError):
                fixed_cents = tl_to_cents(str(fixed_raw))

    percent_rate_raw = _row_value(finans_row, "sozlesme_yuzdesi")
    try:
        percent_rate = Decimal(str(percent_rate_raw))
    except Exception:
        percent_rate = Decimal("0")

    target_cents = int(_row_value(finans_row, "tahsil_hedef_cents") or 0)
    percent_cents = 0
    if percent_rate and target_cents:
        percent_cents = int(
            (Decimal(target_cents) * percent_rate / Decimal("100")).quantize(
                Decimal("1"), rounding=ROUND_HALF_UP
            )
        )

    plan_total = fixed_cents
    if percent_flag == 0:
        plan_total += percent_cents

    if plan_total <= 0:
        return []
    start_date = baslangic_tarihi or date.today()
    amounts: List[int] = []
    base = plan_total // taksit_sayisi
    remainder = plan_total - (base * taksit_sayisi)
    for index in range(taksit_sayisi):
        value = base
        if index == taksit_sayisi - 1:
            value += remainder
        amounts.append(value)

    installments: List[Dict[str, Any]] = []
    period = (periyot or "Ay").lower()
    if period == "ay":
        normalized_day = max(1, min(28, int(vade_gunu or 1)))
        for idx, amount in enumerate(amounts):
            if idx == 0:
                due_date = start_date
            else:
                due_date = _add_months(start_date, idx, normalized_day)
            installments.append(
                {
                    "sira": idx + 1,
                    "vade_tarihi": due_date.strftime("%Y-%m-%d"),
                    "tutar_cents": amount,
                    "durum": "Ödenecek",
                    "odeme_tarihi": None,
                    "aciklama": "",
                }
            )
    else:
        for idx, amount in enumerate(amounts):
            due_date = start_date + timedelta(days=7 * idx)
            installments.append(
                {
                    "sira": idx + 1,
                    "vade_tarihi": due_date.strftime("%Y-%m-%d"),
                    "tutar_cents": amount,
                    "durum": "Ödenecek",
                    "odeme_tarihi": None,
                    "aciklama": "",
                }
            )
    return installments


def get_finans_owner_dosya(finans_id: int) -> int:
    conn = get_connection()
    cur = conn.cursor()
    cur.execute("SELECT dosya_id FROM finans WHERE id = ?", (finans_id,))
    row = cur.fetchone()
    conn.close()
    return int(row[0]) if row else 0


def get_payment_plan(finans_id: int) -> tuple[Dict[str, Any] | None, List[Dict[str, Any]]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM odeme_plani WHERE finans_id = ?", (finans_id,))
    plan_row = cur.fetchone()
    cur.execute(
        "SELECT * FROM taksitler WHERE finans_id = ? ORDER BY vade_tarihi",
        (finans_id,),
    )
    taksitler = [dict(row) for row in cur.fetchall()]
    changed = _apply_overdue_statuses(taksitler, cursor=cur)
    if changed:
        conn.commit()
    conn.close()
    plan = dict(plan_row) if plan_row else None
    return plan, taksitler


def save_payment_plan(
    finans_id: int,
    plan_data: Dict[str, Any],
    taksitler: List[Dict[str, Any]],
    *,
    sync_payments: bool = True,
) -> None:
    conn = get_connection()
    try:
        cur = conn.cursor()
        if plan_data:
            cur.execute(
                """
                INSERT INTO odeme_plani (finans_id, taksit_sayisi, periyot, vade_gunu, baslangic_tarihi, aciklama)
                VALUES (?, ?, ?, ?, ?, ?)
                ON CONFLICT(finans_id) DO UPDATE SET
                    taksit_sayisi=excluded.taksit_sayisi,
                    periyot=excluded.periyot,
                    vade_gunu=excluded.vade_gunu,
                    baslangic_tarihi=excluded.baslangic_tarihi,
                    aciklama=excluded.aciklama
                """,
                (
                    finans_id,
                    int(plan_data.get("taksit_sayisi") or 0),
                    plan_data.get("periyot") or "Ay",
                    int(plan_data.get("vade_gunu") or 7),
                    plan_data.get("baslangic_tarihi"),
                    plan_data.get("aciklama"),
                ),
            )
        else:
            cur.execute("DELETE FROM odeme_plani WHERE finans_id = ?", (finans_id,))

        cur.execute("SELECT id FROM taksitler WHERE finans_id = ?", (finans_id,))
        existing_ids = {row[0] for row in cur.fetchall()}
        kept_ids: Set[int] = set()
        updated_rows: List[Dict[str, Any]] = []

        today = date.today()
        for taksit in taksitler:
            status = _normalize_installment_status(taksit, today=today)
            installment_id = taksit.get("id")
            values = (
                taksit.get("vade_tarihi"),
                int(taksit.get("tutar_cents") or 0),
                status,
                taksit.get("odeme_tarihi"),
                taksit.get("aciklama"),
            )
            if installment_id and installment_id in existing_ids:
                cur.execute(
                    """
                    UPDATE taksitler
                    SET vade_tarihi = ?,
                        tutar_cents = ?,
                        durum = ?,
                        odeme_tarihi = ?,
                        aciklama = ?
                    WHERE id = ? AND finans_id = ?
                    """,
                    (*values, installment_id, finans_id),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO taksitler (finans_id, vade_tarihi, tutar_cents, durum, odeme_tarihi, aciklama)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (finans_id, *values),
                )
                installment_id = cur.lastrowid
            if installment_id:
                kept_ids.add(int(installment_id))
                taksit["id"] = int(installment_id)
            updated_rows.append(taksit)

        removed_ids = existing_ids - kept_ids
        if removed_ids:
            cur.executemany(
                "DELETE FROM taksitler WHERE id = ?",
                ((inst_id,) for inst_id in removed_ids),
            )
            cur.executemany(
                "DELETE FROM odeme_kayitlari WHERE taksit_id = ?",
                ((inst_id,) for inst_id in removed_ids),
            )

        if sync_payments and updated_rows:
            for row in updated_rows:
                status_value = (row.get("durum") or "").strip().casefold()
                if status_value != "ödendi" or not row.get("id"):
                    continue
                try:
                    insert_payment_from_installment(conn, finans_id, row)
                except Exception:
                    # Otomatik ödeme oluşturma hataları plan kaydını engellemesin.
                    pass

        _touch_finans(cur, finans_id)
        recalculate_finans_totals(finans_id, cur)
        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Ödeme planı kaydedilemedi.") from exc
    finally:
        conn.close()


def reset_payment_plan(finans_id: int, keep_paid: bool = False) -> int:
    """Ödeme planını sıfırla.

    Args:
        finans_id: Finans kaydının ID'si
        keep_paid: True ise sadece ödenmemiş taksitleri sil, False ise tüm planı sil

    Returns:
        Silinen taksit sayısı
    """
    conn = get_connection()
    try:
        cur = conn.cursor()
        deleted_count = 0

        if keep_paid:
            # Sadece ödenmemiş taksitleri sil (Ödendi olmayanlar)
            cur.execute(
                """
                SELECT id FROM taksitler
                WHERE finans_id = ? AND (durum IS NULL OR TRIM(durum) != 'Ödendi')
                """,
                (finans_id,),
            )
            unpaid_ids = [row[0] for row in cur.fetchall()]
            deleted_count = len(unpaid_ids)

            if unpaid_ids:
                cur.executemany(
                    "DELETE FROM odeme_kayitlari WHERE taksit_id = ?",
                    ((tid,) for tid in unpaid_ids),
                )
                cur.executemany(
                    "DELETE FROM taksitler WHERE id = ?",
                    ((tid,) for tid in unpaid_ids),
                )

            # Kalan taksit sayısını güncelle
            cur.execute(
                "SELECT COUNT(*) FROM taksitler WHERE finans_id = ?",
                (finans_id,),
            )
            remaining = cur.fetchone()[0]
            cur.execute(
                "UPDATE odeme_plani SET taksit_sayisi = ? WHERE finans_id = ?",
                (remaining, finans_id),
            )
        else:
            # Tüm planı sil
            cur.execute(
                "SELECT COUNT(*) FROM taksitler WHERE finans_id = ?",
                (finans_id,),
            )
            deleted_count = cur.fetchone()[0]

            cur.execute(
                "DELETE FROM odeme_kayitlari WHERE finans_id = ? AND taksit_id IS NOT NULL",
                (finans_id,),
            )
            cur.execute("DELETE FROM taksitler WHERE finans_id = ?", (finans_id,))
            cur.execute("DELETE FROM odeme_plani WHERE finans_id = ?", (finans_id,))

        _touch_finans(cur, finans_id)
        recalculate_finans_totals(finans_id, cur)
        conn.commit()
        return deleted_count
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Ödeme planı sıfırlanamadı.") from exc
    finally:
        conn.close()


def harici_reset_payment_plan(conn, harici_finans_id: int, keep_paid: bool = False) -> int:
    """Harici finans ödeme planını sıfırla.

    Args:
        conn: Veritabanı bağlantısı
        harici_finans_id: Harici finans kaydının ID'si
        keep_paid: True ise sadece ödenmemiş taksitleri sil, False ise tüm planı sil

    Returns:
        Silinen taksit sayısı
    """
    cur = conn.cursor()
    deleted_count = 0

    if keep_paid:
        # Sadece ödenmemiş taksitleri sil (Ödendi olmayanlar)
        cur.execute(
            """
            SELECT id FROM odeme_plani_harici
            WHERE harici_finans_id = ? AND (durum IS NULL OR TRIM(durum) != 'Ödendi')
            """,
            (harici_finans_id,),
        )
        unpaid_ids = [row[0] for row in cur.fetchall()]
        deleted_count = len(unpaid_ids)

        if unpaid_ids:
            cur.executemany(
                "DELETE FROM odemeler_harici WHERE plan_taksit_id = ?",
                ((tid,) for tid in unpaid_ids),
            )
            cur.executemany(
                "DELETE FROM odeme_plani_harici WHERE id = ?",
                ((tid,) for tid in unpaid_ids),
            )
    else:
        # Tüm planı sil
        cur.execute(
            "SELECT COUNT(*) FROM odeme_plani_harici WHERE harici_finans_id = ?",
            (harici_finans_id,),
        )
        deleted_count = cur.fetchone()[0]

        cur.execute(
            "DELETE FROM odemeler_harici WHERE harici_finans_id = ? AND plan_taksit_id IS NOT NULL",
            (harici_finans_id,),
        )
        cur.execute(
            "DELETE FROM odeme_plani_harici WHERE harici_finans_id = ?",
            (harici_finans_id,),
        )

    harici_recalculate_totals(conn, harici_finans_id)
    conn.commit()
    return deleted_count


def get_payments(finans_id: int) -> List[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        "SELECT * FROM odeme_kayitlari WHERE finans_id = ? ORDER BY tarih",
        (finans_id,),
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def save_payments(finans_id: int, payments: List[Dict[str, Any]]) -> None:
    conn = get_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM odeme_kayitlari WHERE finans_id = ?", (finans_id,))
        for payment in payments:
            cur.execute(
                """
                INSERT INTO odeme_kayitlari (finans_id, tarih, tutar_cents, yontem, aciklama, taksit_id)
                VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    finans_id,
                    payment.get("tarih"),
                    int(payment.get("tutar_cents") or 0),
                    payment.get("yontem"),
                    payment.get("aciklama"),
                    payment.get("taksit_id"),
                ),
            )
        recalculate_finans_totals(finans_id, cur)
        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Ödeme kayıtları kaydedilemedi.") from exc
    finally:
        conn.close()


def get_expenses(finans_id: int) -> List[Dict[str, Any]]:
    conn = get_connection()
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute(
        "SELECT * FROM masraflar WHERE finans_id = ? ORDER BY tarih",
        (finans_id,),
    )
    rows = [dict(row) for row in cur.fetchall()]
    conn.close()
    return rows


def save_expenses(finans_id: int, expenses: List[Dict[str, Any]]) -> None:
    conn = get_connection()
    try:
        cur = conn.cursor()

        # dosya_id'yi al (kasa işlemleri için)
        cur.execute("SELECT dosya_id FROM finans WHERE id = ?", (finans_id,))
        row = cur.fetchone()
        dosya_id = int(row[0]) if row else None

        # Eski masrafları sil
        cur.execute("DELETE FROM masraflar WHERE finans_id = ?", (finans_id,))

        # Eski otomatik kasa girişlerini sil (masraflardan oluşturulanlar)
        if dosya_id:
            cur.execute(
                "DELETE FROM muvekkil_kasasi WHERE dosya_id = ? AND islem_turu = 'Kullanılan Avans'",
                (dosya_id,),
            )

        for expense in expenses:
            tutar_cents = int(expense.get("tutar_cents") or 0)
            odeme_kaynagi = expense.get("odeme_kaynagi") or "Büro"
            tarih = expense.get("tarih")
            kalem = expense.get("kalem") or ""

            cur.execute(
                """
                INSERT INTO masraflar (finans_id, kalem, tutar_cents, tarih, odeme_kaynagi, tahsil_durumu, tahsil_tarihi, aciklama)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    finans_id,
                    kalem,
                    tutar_cents,
                    tarih,
                    odeme_kaynagi,
                    expense.get("tahsil_durumu") or "Bekliyor",
                    expense.get("tahsil_tarihi"),
                    expense.get("aciklama"),
                ),
            )

            # Kasadan yapılan masrafları müvekkil kasasına "Kullanılan Avans" olarak ekle
            if dosya_id and odeme_kaynagi == "Kasadan" and tutar_cents > 0:
                cur.execute(
                    """
                    INSERT INTO muvekkil_kasasi (dosya_id, tarih, tutar_kurus, islem_turu, aciklama)
                    VALUES (?, ?, ?, 'Kullanılan Avans', ?)
                    """,
                    (dosya_id, tarih or "", tutar_cents, f"Masraf: {kalem}"),
                )

        recalculate_finans_totals(finans_id, cur)
        conn.commit()
    except sqlite3.Error as exc:
        conn.rollback()
        raise RuntimeError("Masraf kayıtları kaydedilemedi.") from exc
    finally:
        conn.close()
