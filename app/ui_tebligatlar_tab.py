# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import date, timedelta
from typing import Any, Dict, Optional

from PyQt6.QtCore import Qt, QSortFilterProxyModel, QSettings, QTimer
from PyQt6.QtGui import QKeySequence, QShortcut, QColor
from PyQt6.QtWidgets import (
    QApplication,
    QAbstractItemView,
    QCheckBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QHBoxLayout,
    QHeaderView,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMenu,
    QMessageBox,
    QPushButton,
    QTableView,
    QToolButton,
    QVBoxLayout,
    QWidget,
    QGroupBox,
    QFrame,
)

try:  # pragma: no cover - runtime import guard
    from app.models import get_tebligatlar_list, log_action
    from app.services.tebligat_service import mark_tebligat_complete
except ModuleNotFoundError:  # pragma: no cover
    from models import get_tebligatlar_list, log_action
    from services.tebligat_service import mark_tebligat_complete

try:  # pragma: no cover - runtime import guard
    from app.ui_tebligatlar_model import TebligatlarTableModel, COL_TAMAMLANDI, COL_SIRA
except ModuleNotFoundError:  # pragma: no cover
    from ui_tebligatlar_model import TebligatlarTableModel, COL_TAMAMLANDI, COL_SIRA

try:  # pragma: no cover - runtime import guard
    from app.ui_tebligat_dialog import TebligatDialog
except ModuleNotFoundError:  # pragma: no cover
    from ui_tebligat_dialog import TebligatDialog


class TebligatlarTab(QWidget):
    def __init__(self, current_user: Optional[Dict[str, Any]] = None, parent: Optional[QWidget] = None) -> None:
        super().__init__(parent)
        self.current_user = current_user or {}
        self.user_id = self.current_user.get("id")
        self.table_model = TebligatlarTableModel(self)
        self.proxy_model = QSortFilterProxyModel(self)
        self.proxy_model.setSourceModel(self.table_model)
        self.proxy_model.setSortRole(Qt.ItemDataRole.UserRole)
        self.proxy_model.setDynamicSortFilter(True)
        self.proxy_model.setFilterCaseSensitivity(Qt.CaseSensitivity.CaseInsensitive)
        self.proxy_model.setFilterKeyColumn(-1)  # Tüm sütunlarda ara

        self._all_records: list[Dict[str, Any]] = []
        self._alert_filter: tuple[date, date] | None = None
        self._status_filter: str = "all"
        self._current_filter: str = "all"

        self._build_ui()
        self.load_tebligatlar()

    def _build_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setSpacing(8)

        # Filtre ve arama satırı
        filter_layout = QHBoxLayout()
        filter_layout.setSpacing(8)

        # Tarih filtreleri
        self.filter_buttons: dict[str, QToolButton] = {}
        for key, label in (
            ("all", "Tümü"),
            ("today", "Bugün"),
            ("week", "Bu Hafta"),
            ("month", "Bu Ay"),
            ("overdue", "Gecikmiş"),
        ):
            btn = QToolButton()
            btn.setText(label)
            btn.setCheckable(True)
            btn.setMinimumWidth(70)
            btn.clicked.connect(lambda checked, k=key: self._on_filter_changed(k))
            self.filter_buttons[key] = btn
            filter_layout.addWidget(btn)
        self.filter_buttons["all"].setChecked(True)

        filter_layout.addSpacing(16)

        # Arama kutusu
        filter_layout.addWidget(QLabel("Ara:"))
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("Dosya no, kurum veya içerik ara...")
        self.search_edit.setMinimumWidth(200)
        self.search_edit.textChanged.connect(self._on_search_changed)
        filter_layout.addWidget(self.search_edit)

        filter_layout.addStretch()

        # Export butonları
        self.export_excel_button = QToolButton()
        self.export_excel_button.setText("Excel")
        self.export_excel_button.setToolTip("Excel olarak dışa aktar")
        self.export_excel_button.clicked.connect(self._export_to_excel)
        filter_layout.addWidget(self.export_excel_button)

        self.export_word_button = QToolButton()
        self.export_word_button.setText("Word")
        self.export_word_button.setToolTip("Word formatında dışa aktar (PDF olarak yazdırılabilir)")
        self.export_word_button.clicked.connect(self._export_to_docx)
        filter_layout.addWidget(self.export_word_button)

        # Yeni tebligat butonu
        self.new_button = QPushButton("+ Yeni Tebligat")
        self.new_button.setMinimumWidth(120)
        self.new_button.clicked.connect(self.add_tebligat)
        filter_layout.addWidget(self.new_button)

        # Yenile butonu
        self.refresh_button = QToolButton()
        self.refresh_button.setText("Yenile")
        self.refresh_button.clicked.connect(self.load_tebligatlar)
        filter_layout.addWidget(self.refresh_button)

        layout.addLayout(filter_layout)

        self.table_view = QTableView(self)
        self.table_view.setModel(self.proxy_model)
        self.table_view.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table_view.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.table_view.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table_view.setWordWrap(False)
        self.table_view.verticalHeader().setVisible(False)
        self.table_view.setVerticalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.table_view.setHorizontalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.table_view.setSortingEnabled(True)
        self.table_view.doubleClicked.connect(self._open_selected_from_index)
        header = self.table_view.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
        header.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        header.customContextMenuRequested.connect(self._open_header_menu)

        # Kaydedilmiş sütun genişliklerini yükle
        self._load_column_widths()

        # Sütun genişliği değiştiğinde otomatik kaydet
        self._column_save_timer = QTimer(self)
        self._column_save_timer.setSingleShot(True)
        self._column_save_timer.setInterval(500)
        self._column_save_timer.timeout.connect(self._save_column_widths)
        header.sectionResized.connect(self._on_section_resized)

        self._copy_shortcut = QShortcut(
            QKeySequence(QKeySequence.StandardKey.Copy), self.table_view
        )
        self._copy_shortcut.activated.connect(self._copy_selection)

        # Checkbox sütununa tıklama için
        self.table_view.clicked.connect(self._on_cell_clicked)

        layout.addWidget(self.table_view)

    def _on_cell_clicked(self, index) -> None:
        """Hücreye tıklandığında - checkbox sütunu ise tamamlandı durumunu değiştir."""
        source_index = self.proxy_model.mapToSource(index)
        if source_index.column() == COL_TAMAMLANDI:
            # ID'yi SIRA sütunundan al (UserRole)
            id_index = self.table_model.index(source_index.row(), COL_SIRA)
            tebligat_id = self.table_model.data(id_index, Qt.ItemDataRole.UserRole)
            if tebligat_id:
                # Mevcut durumu al
                current_status = self.table_model.data(source_index, Qt.ItemDataRole.UserRole)
                new_status = not bool(current_status)
                try:
                    mark_tebligat_complete(int(tebligat_id), new_status)
                    self.load_tebligatlar()
                except Exception as exc:
                    QMessageBox.warning(self, "Hata", f"İşlem tamamlanamadı:\n{exc}")

    def _on_filter_changed(self, key: str) -> None:
        """Filtre butonu değiştiğinde."""
        for name, btn in self.filter_buttons.items():
            btn.setChecked(name == key)
        self._current_filter = key
        self._apply_filters()

    def _on_search_changed(self, text: str) -> None:
        """Arama metni değiştiğinde."""
        self.proxy_model.setFilterFixedString(text)

    def _apply_filters(self) -> None:
        """Tarih ve durum filtrelerini uygula."""
        today = date.today()
        filtered = list(self._all_records)

        if self._current_filter == "today":
            filtered = [r for r in filtered if self._coerce_date(r.get("is_son_gunu")) == today]
        elif self._current_filter == "week":
            week_end = today + timedelta(days=7)
            filtered = [r for r in filtered if self._is_in_range(r, today, week_end)]
        elif self._current_filter == "month":
            month_end = today + timedelta(days=30)
            filtered = [r for r in filtered if self._is_in_range(r, today, month_end)]
        elif self._current_filter == "overdue":
            filtered = [r for r in filtered if self._is_overdue(r)]

        self.table_model.set_records(filtered)
        self.proxy_model.invalidate()

    def _is_in_range(self, record: Dict[str, Any], start: date, end: date) -> bool:
        """Kaydın tarihinin aralıkta olup olmadığını kontrol et."""
        is_son_gunu = self._coerce_date(record.get("is_son_gunu"))
        if is_son_gunu and start <= is_son_gunu <= end:
            return True
        return False

    def _is_overdue(self, record: Dict[str, Any]) -> bool:
        """Kaydın gecikmiş olup olmadığını kontrol et (sadece tamamlanmamış olanlar)."""
        # Tamamlanmış kayıtlar gecikmiş sayılmaz
        if record.get("tamamlandi"):
            return False
        is_son_gunu = self._coerce_date(record.get("is_son_gunu"))
        if is_son_gunu and is_son_gunu < date.today():
            return True
        return False

    def load_tebligatlar(self) -> None:
        try:
            records = get_tebligatlar_list()
        except Exception as exc:  # pragma: no cover - GUI güvenliği
            QMessageBox.critical(self, "Hata", f"Tebligatlar yüklenemedi:\n{exc}")
            return
        self._all_records = list(records)
        self._apply_filters()

    def add_tebligat(self) -> None:
        dialog = TebligatDialog(self, current_user=self.current_user)
        if dialog.exec() == int(dialog.DialogCode.Accepted):
            if dialog.last_action == "insert" and self.user_id:
                log_action(self.user_id, "add_tebligat", dialog.tebligat_id)
            self.load_tebligatlar()

    def _open_selected_from_index(self, index) -> None:
        if not index.isValid():
            return
        source_index = self.proxy_model.mapToSource(index)
        record = self.table_model.record_at(source_index.row())
        if not record:
            return
        self._open_tebligat(record)

    def _open_tebligat(self, record: Dict[str, Any]) -> None:
        tebligat_id = record.get("id")
        if tebligat_id is None:
            return
        dialog = TebligatDialog(
            self,
            tebligat_id=int(tebligat_id),
            current_user=self.current_user,
        )
        if dialog.exec() != int(dialog.DialogCode.Accepted):
            return
        action = dialog.last_action
        if action == "delete" and self.user_id:
            log_action(self.user_id, "delete_tebligat", int(tebligat_id))
        elif action == "update" and self.user_id:
            log_action(self.user_id, "update_tebligat", int(tebligat_id))
        self.load_tebligatlar()

    def _open_header_menu(self, pos) -> None:
        header = self.table_view.horizontalHeader()
        if header is None:
            return
        menu = QMenu(self.table_view)
        for visual in range(header.count()):
            logical = header.logicalIndex(visual)
            title = self.table_model.headerData(
                logical,
                Qt.Orientation.Horizontal,
                Qt.ItemDataRole.DisplayRole,
            )
            label = str(title) if title else f"Sütun {logical + 1}"
            action = menu.addAction(label)
            action.setCheckable(True)
            action.setChecked(not self.table_view.isColumnHidden(logical))
            action.toggled.connect(
                lambda checked, col=logical: self.table_view.setColumnHidden(col, not checked)
            )
        menu.exec(header.mapToGlobal(pos))

    def _copy_selection(self) -> None:
        selection = self.table_view.selectionModel()
        if selection is None or not selection.hasSelection():
            return
        indexes = selection.selectedIndexes()
        if not indexes:
            return
        sorted_indexes = sorted(indexes, key=lambda idx: (idx.row(), idx.column()))
        rows: dict[int, list] = {}
        for index in sorted_indexes:
            rows.setdefault(index.row(), []).append(index)
        text_rows: list[str] = []
        for row in sorted(rows):
            columns = sorted(rows[row], key=lambda idx: idx.column())
            values = []
            for index in columns:
                data = index.data(Qt.ItemDataRole.DisplayRole)
                values.append(str(data) if data is not None else "")
            text_rows.append("\t".join(values))
        QApplication.clipboard().setText("\n".join(text_rows))

    def apply_alert_date_filter(self, start: date, end: date) -> None:
        if end < start:
            start, end = end, start
        self._alert_filter = (start, end)
        self._apply_alert_filter()

    def clear_alert_date_filter(self) -> None:
        self._alert_filter = None
        self._apply_alert_filter()

    def _apply_alert_filter(self) -> None:
        rows = list(self._all_records)
        if self._alert_filter is not None:
            start, end = self._alert_filter
            rows = [row for row in rows if self._matches_alert_range(row, start, end)]
        self.table_model.set_records(rows)
        self.proxy_model.invalidate()

    def _matches_alert_range(self, record: Dict[str, Any], start: date, end: date) -> bool:
        for key in ("is_son_gunu", "teblig_tarihi", "geldigi_tarih", "tebligat_tarihi"):
            row_date = self._coerce_date(record.get(key))
            if row_date is not None and start <= row_date <= end:
                return True
        return False

    @staticmethod
    def _coerce_date(value: Any) -> date | None:
        if value is None:
            return None
        if isinstance(value, date):
            return value
        text = str(value).strip()
        if not text:
            return None
        try:
            return date.fromisoformat(text[:10])
        except ValueError:
            return None

    def _get_visible_records(self) -> list[Dict[str, Any]]:
        """Proxy model üzerinden görünen kayıtları döndürür."""
        records = []
        for row in range(self.proxy_model.rowCount()):
            source_index = self.proxy_model.mapToSource(self.proxy_model.index(row, 0))
            record = self.table_model.record_at(source_index.row())
            if record:
                records.append(record)
        return records

    def _export_to_excel(self) -> None:
        """Tebligatları Excel formatında dışa aktar."""
        records = self._get_visible_records()
        if not records:
            QMessageBox.information(self, "Bilgi", "Dışa aktarılacak kayıt bulunamadı.")
            return

        # Sütun seçim dialogu
        dialog = TebligatExportDialog(records, self)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        selected_records = dialog.get_selected_records()
        if not selected_records:
            return

        # Dosya kaydetme dialogu
        path, _ = QFileDialog.getSaveFileName(
            self, "Excel Olarak Kaydet", "tebligatlar.xlsx", "Excel Dosyası (*.xlsx)"
        )
        if not path:
            return

        try:
            self._write_excel(path, selected_records)
            QMessageBox.information(self, "Başarılı", f"Dosya kaydedildi:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Hata", f"Dosya kaydedilemedi:\n{exc}")

    def _export_to_docx(self) -> None:
        """Tebligatları Word formatında dışa aktar."""
        records = self._get_visible_records()
        if not records:
            QMessageBox.information(self, "Bilgi", "Dışa aktarılacak kayıt bulunamadı.")
            return

        # Sütun seçim dialogu
        dialog = TebligatExportDialog(records, self)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        selected_records = dialog.get_selected_records()
        if not selected_records:
            return

        # Dosya kaydetme dialogu
        path, _ = QFileDialog.getSaveFileName(
            self, "Word Olarak Kaydet", "tebligatlar.docx", "Word Dosyası (*.docx)"
        )
        if not path:
            return

        try:
            self._write_docx(path, selected_records)
            QMessageBox.information(self, "Başarılı", f"Dosya kaydedildi:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Hata", f"Dosya kaydedilemedi:\n{exc}")

    def _write_excel(self, path: str, records: list[Dict[str, Any]]) -> None:
        """Tebligatları Excel dosyasına yazar."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl modülü yüklü değil")

        wb = Workbook()
        ws = wb.active
        ws.title = "Tebligatlar"

        headers = ["Sıra", "Dosya No", "Kurum", "Geldiği Tarih", "Tebliğ Tarihi", "İşin Son Günü", "İçerik"]

        # Başlık stilleri
        header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin")
        )

        # Başlık satırı
        ws.append(headers)
        for col_idx, _ in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border

        # Veri satırları
        data_font = Font(name="Calibri", size=10)
        data_alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

        for row_idx, record in enumerate(records, 2):
            row_data = [
                row_idx - 1,
                record.get("dosya_no", "") or "",
                record.get("kurum", "") or "",
                self._format_date_for_export(record.get("geldigi_tarih")),
                self._format_date_for_export(record.get("teblig_tarihi")),
                self._format_date_for_export(record.get("is_son_gunu")),
                record.get("icerik", "") or "",
            ]
            ws.append(row_data)

            for col_idx, _ in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.font = data_font
                cell.alignment = data_alignment
                cell.border = thin_border

        # Sütun genişlikleri
        column_widths = [8, 20, 25, 14, 14, 14, 40]
        for col_idx, width in enumerate(column_widths, 1):
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        # Satır yükseklikleri
        ws.row_dimensions[1].height = 25
        for row_idx in range(2, len(records) + 2):
            ws.row_dimensions[row_idx].height = 20

        ws.auto_filter.ref = ws.dimensions
        ws.page_setup.orientation = "landscape"
        ws.page_setup.fitToPage = True
        ws.page_setup.fitToWidth = 1

        wb.save(path)

    def _write_docx(self, path: str, records: list[Dict[str, Any]]) -> None:
        """Tebligatları Word dosyasına yazar."""
        try:
            from docx import Document
            from docx.shared import Pt, Mm
            from docx.enum.section import WD_ORIENT
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import RGBColor
        except ImportError:
            raise ImportError("python-docx modülü yüklü değil")

        doc = Document()

        # Sayfa ayarları
        for section in doc.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            new_width, new_height = section.page_height, section.page_width
            section.page_width, section.page_height = new_width, new_height
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)
            section.top_margin = Mm(15)
            section.bottom_margin = Mm(15)

        # Başlık
        title = doc.add_paragraph()
        title_run = title.add_run("Tebligat Listesi")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = "Calibri"
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Tarih
        from datetime import datetime
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Oluşturulma Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_run.font.size = Pt(9)
        date_run.font.name = "Calibri"
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        headers = ["Sıra", "Dosya No", "Kurum", "Geldiği Tarih", "Tebliğ Tarihi", "İşin Son Günü", "İçerik"]

        # Tablo
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Başlık hücreleri
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            cell = header_cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.text = ""
            run = paragraph.add_run(str(header))
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            # Başlık arka plan rengi
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = OxmlElement("w:shd")
            shading.set(qn("w:val"), "clear")
            shading.set(qn("w:fill"), "2c3e50")
            tcPr.append(shading)
            run.font.color.rgb = RGBColor(255, 255, 255)

        # Veri satırları
        for row_idx, record in enumerate(records, 1):
            row_cells = table.add_row().cells
            row_data = [
                str(row_idx),
                record.get("dosya_no", "") or "",
                record.get("kurum", "") or "",
                self._format_date_for_export(record.get("geldigi_tarih")),
                self._format_date_for_export(record.get("teblig_tarihi")),
                self._format_date_for_export(record.get("is_son_gunu")),
                record.get("icerik", "") or "",
            ]
            for col_idx, value in enumerate(row_data):
                cell = row_cells[col_idx]
                paragraph = cell.paragraphs[0]
                paragraph.text = ""
                run = paragraph.add_run(str(value))
                run.font.size = Pt(9)
                run.font.name = "Calibri"

        # Sütun genişlikleri
        column_widths_mm = [12, 30, 35, 22, 22, 22, 60]
        for col_index, width_mm in enumerate(column_widths_mm):
            table.columns[col_index].width = Mm(width_mm)
            for row in table.rows:
                row.cells[col_index].width = Mm(width_mm)

        # Footer
        footer = doc.add_paragraph()
        footer_run = footer.add_run(f"\nToplam: {len(records)} kayıt")
        footer_run.font.size = Pt(9)
        footer_run.font.name = "Calibri"
        footer_run.italic = True

        doc.save(path)

    @staticmethod
    def _format_date_for_export(value: Any) -> str:
        """Tarihi export için formatlar."""
        if not value:
            return ""
        try:
            from utils import iso_to_tr
            return iso_to_tr(str(value))
        except Exception:
            return str(value)

    def _on_section_resized(self, logical_index: int, old_size: int, new_size: int) -> None:
        """Sütun genişliği değiştiğinde timer'ı başlat."""
        if hasattr(self, "_column_save_timer"):
            self._column_save_timer.start()

    def _load_column_widths(self) -> None:
        """Kaydedilmiş sütun genişliklerini yükle."""
        header = self.table_view.horizontalHeader()
        if header is None:
            return
        try:
            settings = QSettings("MyCompany", "LexTakip")
            stored = settings.value("tebligatlar/col_widths")
            widths: list[int] = []
            if isinstance(stored, (list, tuple)):
                for value in stored:
                    try:
                        widths.append(int(value))
                    except (TypeError, ValueError):
                        continue
            elif isinstance(stored, str):
                for part in stored.split(","):
                    part = part.strip()
                    if part:
                        try:
                            widths.append(int(part))
                        except ValueError:
                            continue
            if widths:
                for col, width in enumerate(widths):
                    if 0 <= col < header.count() and width > 0:
                        header.resizeSection(col, width)
        except Exception:
            pass

    def _save_column_widths(self) -> None:
        """Sütun genişliklerini kaydet."""
        header = self.table_view.horizontalHeader()
        if header is None:
            return
        try:
            widths = [header.sectionSize(col) for col in range(header.count())]
            settings = QSettings("MyCompany", "LexTakip")
            settings.setValue("tebligatlar/col_widths", widths)
            settings.sync()
        except Exception:
            pass


class TebligatExportDialog(QDialog):
    """Tebligat export seçim dialogu."""

    def __init__(self, records: list[Dict[str, Any]], parent=None):
        super().__init__(parent)
        self.setWindowTitle("Dışa Aktarılacak Kayıtları Seç")
        self.setMinimumWidth(500)
        self.setMinimumHeight(400)

        self._records = records
        self._selected_indices: list[int] = []

        layout = QVBoxLayout(self)

        # Bilgi etiketi
        info_label = QLabel(f"Toplam {len(records)} kayıt bulundu. Dışa aktarmak istediklerinizi seçin:")
        layout.addWidget(info_label)

        # Aralık seçimi
        range_group = QGroupBox("Aralık Seçimi")
        range_layout = QHBoxLayout(range_group)
        range_layout.addWidget(QLabel("Aralık:"))
        self.range_edit = QLineEdit()
        self.range_edit.setPlaceholderText("Örn: 3-6 veya 1,3,5-7")
        self.range_edit.setMinimumWidth(150)
        self.range_edit.returnPressed.connect(self._apply_range)
        range_layout.addWidget(self.range_edit)
        self.apply_range_btn = QPushButton("Uygula")
        self.apply_range_btn.clicked.connect(self._apply_range)
        range_layout.addWidget(self.apply_range_btn)
        range_layout.addStretch()
        layout.addWidget(range_group)

        # Seçim butonları
        select_layout = QHBoxLayout()
        self.select_all_btn = QPushButton("Tümünü Seç")
        self.select_all_btn.clicked.connect(self._select_all)
        select_layout.addWidget(self.select_all_btn)

        self.deselect_all_btn = QPushButton("Hiçbirini Seçme")
        self.deselect_all_btn.clicked.connect(self._deselect_all)
        select_layout.addWidget(self.deselect_all_btn)
        select_layout.addStretch()
        layout.addLayout(select_layout)

        # Kayıt listesi
        self.list_widget = QListWidget()
        for idx, record in enumerate(records):
            dosya_no = record.get("dosya_no", "") or "-"
            kurum = record.get("kurum", "") or "-"
            is_son_gunu = self._format_date_display(record.get("is_son_gunu"))

            item_text = f"{idx + 1}. {dosya_no} - {kurum} (Son gün: {is_son_gunu})"
            item = QListWidgetItem(item_text)
            item.setCheckState(Qt.CheckState.Checked)
            item.setData(Qt.ItemDataRole.UserRole, idx)
            self.list_widget.addItem(item)

        layout.addWidget(self.list_widget)

        # Dialog butonları
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)

    def _select_all(self) -> None:
        for i in range(self.list_widget.count()):
            self.list_widget.item(i).setCheckState(Qt.CheckState.Checked)

    def _deselect_all(self) -> None:
        for i in range(self.list_widget.count()):
            self.list_widget.item(i).setCheckState(Qt.CheckState.Unchecked)

    def _apply_range(self) -> None:
        """Aralık metnini parse edip ilgili kayıtları seçer.

        Örnek formatlar:
        - "3-6" -> 3, 4, 5, 6 numaralı kayıtları seçer
        - "1,3,5-7" -> 1, 3, 5, 6, 7 numaralı kayıtları seçer
        - "5" -> sadece 5 numaralı kaydı seçer
        """
        range_text = self.range_edit.text().strip()
        if not range_text:
            return

        # Önce tümünü temizle
        self._deselect_all()

        # Seçilecek indisleri hesapla
        selected_indices: set[int] = set()
        total = self.list_widget.count()

        # Virgülle ayrılmış parçaları işle
        parts = range_text.replace(" ", "").split(",")
        for part in parts:
            if "-" in part:
                # Aralık: "3-6"
                try:
                    range_parts = part.split("-")
                    if len(range_parts) == 2:
                        start = int(range_parts[0])
                        end = int(range_parts[1])
                        if start > end:
                            start, end = end, start
                        for i in range(start, end + 1):
                            if 1 <= i <= total:
                                selected_indices.add(i - 1)  # 0-indexed
                except ValueError:
                    continue
            else:
                # Tek sayı: "5"
                try:
                    num = int(part)
                    if 1 <= num <= total:
                        selected_indices.add(num - 1)  # 0-indexed
                except ValueError:
                    continue

        # Seçimleri uygula
        for idx in selected_indices:
            item = self.list_widget.item(idx)
            if item:
                item.setCheckState(Qt.CheckState.Checked)

        # Bilgi ver
        if selected_indices:
            self.range_edit.setStyleSheet("")
        else:
            self.range_edit.setStyleSheet("border: 1px solid red;")

    def get_selected_records(self) -> list[Dict[str, Any]]:
        """Seçilen kayıtları döndürür."""
        selected = []
        for i in range(self.list_widget.count()):
            item = self.list_widget.item(i)
            if item.checkState() == Qt.CheckState.Checked:
                idx = item.data(Qt.ItemDataRole.UserRole)
                if 0 <= idx < len(self._records):
                    selected.append(self._records[idx])
        return selected

    @staticmethod
    def _format_date_display(value: Any) -> str:
        if not value:
            return "-"
        try:
            from utils import iso_to_tr
            return iso_to_tr(str(value))
        except Exception:
            return str(value)
